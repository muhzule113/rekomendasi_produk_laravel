from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "docs" / "laporan" / "Pengujian_Blackbox_Sesuai_Whitebox.docx"


def set_run_font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=4)
    table.style = "Table Grid"
    headers = ["No", "Tes Faktor", "Hasil", "Keterangan"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, 11, True)
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9D9D9")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            if c in (0, 2):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run_font(run, 11, False)
    widths = [Cm(1.2), Cm(6.5), Cm(1.8), Cm(6.5)]
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


SECTIONS = [
    (
        "1. Daftar Akun User",
        "Acuan white-box: V(G) = 2 | J1: 1-2-3-5-6-7 | J2: 1-2-3-4-2-3-5-6-7",
        [
            (
                "1",
                'Mengisi semua form valid (nama, email unik, no HP, alamat, password ≥ 6, konfirmasi sama) lalu menekan tombol "Daftar Sekarang".',
                "✓",
                'Sistem menerima data, menyimpan user pelanggan/aktif, menampilkan pesan "Pendaftaran berhasil! Silakan login." (Jalur J1)',
            ),
            (
                "2",
                'Mengisi form tidak valid (email sudah terdaftar / konfirmasi berbeda / password < 6) lalu menekan "Daftar Sekarang".',
                "✓",
                "Sistem menolak input, menampilkan pesan error validasi, form dapat diisi ulang. (Jalur J2)",
            ),
        ],
    ),
    (
        "2. Login User",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                'Mengisi email & password pelanggan valid, role pelanggan, menekan "Masuk".',
                "✓",
                "Sistem menerima akses dan mengarahkan ke halaman Katalog Produk (/produk). (Jalur J1)",
            ),
            (
                "2",
                'Mengisi email/password salah atau role tidak sesuai, menekan "Masuk".',
                "✓",
                "Sistem menampilkan pesan error dan tetap di halaman login. (Jalur J2)",
            ),
        ],
    ),
    (
        "3. Login Admin",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                'Mengisi email & password admin valid, role admin, menekan "Masuk".',
                "✓",
                "Sistem menerima akses dan mengarahkan ke Dashboard Admin. (Jalur J1)",
            ),
            (
                "2",
                'Mengisi kredensial salah / role tidak sesuai admin, menekan "Masuk".',
                "✓",
                "Sistem menampilkan pesan error dan tetap di halaman login. (Jalur J2)",
            ),
        ],
    ),
    (
        "4. Menu Rekomendasi Produk User",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login sebagai pelanggan, membuka menu Rekomendasi saat ada kandidat rekomendasi.",
                "✓",
                "Sistem menampilkan daftar rekomendasi produk. (Jalur Ada)",
            ),
            (
                "2",
                "Login sebagai pelanggan, membuka Rekomendasi saat tidak ada kandidat.",
                "✓",
                "Sistem menampilkan kondisi rekomendasi kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "5. Menu Beranda User",
        "Acuan white-box: V(G) = 1",
        [
            (
                "1",
                "Membuka halaman utama / Beranda.",
                "✓",
                "Sistem menampilkan beranda beserta produk unggulan. (Jalur tunggal)",
            ),
        ],
    ),
    (
        "6. Menu Katalog User",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Membuka menu Katalog saat ada produk aktif.",
                "✓",
                "Sistem menampilkan daftar produk katalog. (Jalur Ada)",
            ),
            (
                "2",
                "Membuka Katalog saat tidak ada produk yang dapat ditampilkan.",
                "✓",
                "Sistem menampilkan katalog kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "7. Menu Keranjang User",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Membuka Keranjang yang berisi produk.",
                "✓",
                "Sistem menampilkan daftar produk di keranjang. (Jalur Ada)",
            ),
            (
                "2",
                "Membuka Keranjang dalam keadaan kosong.",
                "✓",
                "Sistem menampilkan keranjang kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "8. Menu Checkout User",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login pelanggan, keranjang berisi item, membuka Checkout, mengisi form valid, menyelesaikan pesanan.",
                "✓",
                "Sistem menyimpan pesanan dan menampilkan konfirmasi sukses. (Jalur Ada)",
            ),
            (
                "2",
                "Login pelanggan, membuka Checkout dengan keranjang kosong.",
                "✓",
                "Sistem menampilkan kondisi keranjang kosong / tidak memproses pesanan. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "9. Menu Riwayat User",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login pelanggan yang punya transaksi, membuka Riwayat.",
                "✓",
                "Sistem menampilkan daftar riwayat transaksi. (Jalur Ada)",
            ),
            (
                "2",
                "Login pelanggan tanpa transaksi, membuka Riwayat.",
                "✓",
                "Sistem menampilkan riwayat kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "10. Admin Dashboard",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Dashboard saat data statistik tersedia.",
                "✓",
                "Sistem menampilkan dashboard statistik. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka Dashboard saat statistik kosong.",
                "✓",
                "Sistem menampilkan dashboard dengan informasi kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "11. Admin Kelola Produk",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Kelola Produk saat ada data produk.",
                "✓",
                "Halaman kelola produk menampilkan daftar produk. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka Kelola Produk saat data produk kosong.",
                "✓",
                "Menampilkan kondisi data produk kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "12. Admin Kelola Transaksi",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Kelola Transaksi saat ada transaksi.",
                "✓",
                "Daftar transaksi ditampilkan. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka Kelola Transaksi saat tidak ada transaksi.",
                "✓",
                "Menampilkan data transaksi kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "13. Admin Data Pelanggan",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Data Pelanggan saat ada pelanggan.",
                "✓",
                "Daftar pelanggan ditampilkan. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka Data Pelanggan saat kosong.",
                "✓",
                "Menampilkan data pelanggan kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "14. Admin Analisis Rekomendasi",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Analisis Rekomendasi saat data analisis ada.",
                "✓",
                "Halaman analisis ditampilkan. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka analisis saat data kosong.",
                "✓",
                "Menampilkan analisis kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "15. Admin Riwayat Upload",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Riwayat Upload saat ada log upload.",
                "✓",
                "Daftar riwayat upload ditampilkan. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka riwayat upload saat kosong.",
                "✓",
                "Menampilkan riwayat kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "16. Admin Ulasan & Rating",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Ulasan & Rating saat ada ulasan.",
                "✓",
                "Daftar ulasan ditampilkan. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka ulasan saat kosong.",
                "✓",
                "Menampilkan ulasan kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "17. Admin Laporan",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, membuka Laporan saat data laporan ada.",
                "✓",
                "Halaman laporan ditampilkan. (Jalur Ada)",
            ),
            (
                "2",
                "Login admin, membuka laporan saat data kosong.",
                "✓",
                "Menampilkan laporan kosong. (Jalur Kosong)",
            ),
        ],
    ),
    (
        "18. Admin Upload Data",
        "Acuan white-box: V(G) = 2",
        [
            (
                "1",
                "Login admin, upload file data valid, proses simpan berhasil.",
                "✓",
                "Sistem menampilkan pesan berhasil. (Jalur Ya)",
            ),
            (
                "2",
                "Login admin, upload file invalid / gagal validasi.",
                "✓",
                "Sistem menampilkan pesan error; dapat mengulang upload. (Jalur Tidak)",
            ),
        ],
    ),
]


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2)
        s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t.add_run("Pengujian Black-Box"), 14, True)
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(t2.add_run("(Diselaraskan dengan Jalur White-Box)"), 12, True)

    for title, note, rows in SECTIONS:
        h = doc.add_paragraph()
        set_run_font(h.add_run(title), 12, True)
        n = doc.add_paragraph()
        set_run_font(n.add_run(note), 11, False)
        add_table(doc, rows)
        ss = doc.add_paragraph()
        set_run_font(ss.add_run("Screenshot:"), 11, False)
        blank = doc.add_paragraph()
        set_run_font(blank.add_run("(sisipkan screenshot di sini)"), 11, False)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
