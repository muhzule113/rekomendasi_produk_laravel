from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "Lampiran_Kode_Program_Inti.docx"


def set_cell_shading(cell, fill="FFFFFF"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:color"), "auto")


def set_section_columns(section, count=1, space_twips=360):
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space_twips))


def set_margins(section):
    # Format 4-4-3-3 cm: atas, kiri, bawah, kanan.
    section.top_margin = Cm(4)
    section.left_margin = Cm(4)
    section.bottom_margin = Cm(3)
    section.right_margin = Cm(3)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for name, size, bold, color in [
        ("Lampiran Title", 14, True, "000000"),
        ("Lampiran Subtitle", 12, False, "404040"),
        ("Code Heading", 11, True, "1F4E79"),
        ("Code Path", 8.5, False, "666666"),
    ]:
        if name in document.styles:
            style = document.styles[name]
        else:
            style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor.from_string(color)


def read_lines(relative_path, start, end):
    path = ROOT / relative_path
    lines = path.read_text(encoding="utf-8").splitlines()
    selected = lines[start - 1 : end]
    return [(start + index, line) for index, line in enumerate(selected)]


def add_code_block(document, title, relative_path, start, end, note=None):
    heading = document.add_paragraph(style="Code Heading")
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(1)
    heading.add_run(title)

    path_p = document.add_paragraph(style="Code Path")
    path_p.paragraph_format.keep_with_next = True
    path_p.paragraph_format.space_after = Pt(2)
    path_p.add_run(f"Sumber: {relative_path}")

    if note:
        note_p = document.add_paragraph()
        note_p.paragraph_format.space_before = Pt(0)
        note_p.paragraph_format.space_after = Pt(3)
        note_p.paragraph_format.keep_with_next = True
        note_run = note_p.add_run(note)
        note_run.italic = True
        note_run.font.name = "Times New Roman"
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(90, 90, 90)

    # Kode ditulis sebagai paragraf biasa, bukan tabel, agar tidak ada
    # style tabel atau shading bawaan yang dapat menampilkan latar abu-abu.
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, (_line_number, line) in enumerate(read_lines(relative_path, start, end)):
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(35, 45, 55)
        if index < (end - start):
            run.add_break()

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)


def add_footer(section):
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Lampiran Kode Program Inti — ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)
    add_page_field(paragraph)


def build_document():
    document = Document()
    configure_styles(document)

    first_section = document.sections[0]
    set_margins(first_section)
    set_section_columns(first_section, 1)

    title = document.add_paragraph(style="Lampiran Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(14)
    title.paragraph_format.space_after = Pt(6)
    title.add_run("LAMPIRAN A\nKODE PROGRAM INTI")

    subtitle = document.add_paragraph(style="Lampiran Subtitle")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.add_run("Sistem Rekomendasi Produk Toko Sinar Manis")

    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    intro.paragraph_format.first_line_indent = Cm(1)
    intro.paragraph_format.line_spacing = 1.15
    intro.add_run(
        "Lampiran ini memuat bagian inti dari project Laravel yang digunakan dalam "
        "implementasi sistem. Kode ditampilkan secara ringkas, meliputi routing, "
        "model produk, pencatatan transaksi, pembentukan matriks transaksi, "
        "perhitungan Item-Based Collaborative Filtering dengan Cosine Similarity, "
        "serta penyajian rekomendasi kepada pelanggan."
    )

    info = document.add_paragraph()
    info.paragraph_format.space_before = Pt(12)
    info.paragraph_format.line_spacing = 1.15
    info.add_run("Keterangan format: ").bold = True
    info.add_run(
        "ukuran kertas A4, margin atas 4 cm, kiri 4 cm, bawah 3 cm, kanan 3 cm, "
        "dan bagian kode menggunakan dua kolom agar kodingan lebih bersih "
        "dan mudah dibaca."
    )

    code_section = document.add_section(WD_SECTION.NEW_PAGE)
    set_margins(code_section)
    set_section_columns(code_section, 2)

    add_code_block(
        document,
        "A. Routing halaman utama dan rekomendasi",
        "routes/web.php",
        8,
        42,
        "Route menghubungkan halaman katalog, keranjang, checkout, rekomendasi, dan area admin dengan controller terkait.",
    )
    add_code_block(
        document,
        "B. Model produk",
        "app/Models/Product.php",
        1,
        46,
        "Model mendefinisikan tabel products, primary key, atribut yang dapat diisi, casting, dan relasi data.",
    )
    add_code_block(
        document,
        "C. Proses checkout dan pencatatan transaksi",
        "app/Http/Controllers/CheckoutController.php",
        45,
        109,
        "Transaksi hanya dibuat setelah isi keranjang dan stok divalidasi; detail transaksi kemudian disimpan ke transaction_items.",
    )
    add_code_block(
        document,
        "D. Pembentukan matriks user-item dan Cosine Similarity",
        "app/Services/RecommenderService.php",
        38,
        162,
        "Data transaksi selesai dan telah dibayar diubah menjadi matriks biner, kemudian dihitung similarity antarpasangan produk.",
    )
    add_code_block(
        document,
        "E. Penyusunan rekomendasi personal",
        "app/Services/RecommenderService.php",
        291,
        390,
        "Produk yang belum dibeli disaring berdasarkan stok dan skor similarity, lalu diurutkan berdasarkan prediction score.",
    )
    add_code_block(
        document,
        "F. Controller halaman rekomendasi",
        "app/Http/Controllers/RekomendasiController.php",
        1,
        63,
        "Controller mengambil hasil rekomendasi, menyesuaikan stok dengan isi keranjang, mencatat log, dan mengirim data ke view.",
    )
    add_code_block(
        document,
        "G. Implementasi Cosine Similarity pada pipeline Python",
        "python/cf/similarity.py",
        1,
        90,
        "Pipeline Python memakai operasi matriks NumPy untuk menghasilkan skor similarity dan menerapkan minimum co-occurrence.",
    )

    for section in document.sections:
        add_footer(section)

    document.core_properties.title = "Lampiran Kode Program Inti"
    document.core_properties.subject = "Lampiran skripsi sistem rekomendasi produk"
    document.core_properties.author = ""
    document.core_properties.comments = "Dibuat dari kode inti project Laravel."
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    output = build_document()
    print(output)
