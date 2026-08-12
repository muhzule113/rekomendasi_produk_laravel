# -*- coding: utf-8 -*-
"""Tambahkan instrumen dan format hasil SUS ke dokumen Bab III–IV."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "laporan" / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_REVISI.docx"
)
OUTPUT = ROOT / "docs" / "laporan" / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_REVISI_SUS.docx"
)


def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_format(paragraph, first_line_indent=True):
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5


def add_heading_at_end(doc, text, level):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        set_run_font(run, size=12 if level == 1 else 11, bold=True)
    return paragraph


def add_body_at_end(doc, text, first_line_indent=True):
    paragraph = doc.add_paragraph()
    set_paragraph_format(paragraph, first_line_indent=first_line_indent)
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_caption_at_end(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, bold=True)
    return paragraph


def add_table_at_end(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    for column_index, header in enumerate(headers):
        cell = table.rows[0].cells[column_index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=9, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            cell = table.rows[row_index].cells[column_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            if column_index in (0, 2, 3, 4, 5, 6):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_run_font(run, size=9)

    if widths_cm:
        for row in table.rows:
            for column_index, width in enumerate(widths_cm):
                row.cells[column_index].width = Cm(width)
        for column_index, width in enumerate(widths_cm):
            table.columns[column_index].width = Cm(width)

    return table


def move_before(element, target):
    """Pindahkan elemen yang baru dibuat sebelum paragraf target."""
    target._p.addprevious(element._p if hasattr(element, "_p") else element._tbl)


def move_table_before(table, target):
    target._p.addprevious(table._tbl)


def move_items_before(items, target):
    for item in items:
        if hasattr(item, "_tbl"):
            move_table_before(item, target)
        else:
            move_before(item, target)


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"Paragraf target tidak ditemukan: {text}")


def main():
    doc = Document(SOURCE)

    # Perjelas rujukan pada paragraf SUS yang sudah ada di Bab III.
    existing = find_paragraph(
        doc,
        "SUS diberikan setelah peserta menyelesaikan tugas relevan (10 pernyataan, skor 0–100). "
        "Ambang interpretasi awal rerata minimal 68 (Bangor et al., 2008). "
        "Skor SUS tidak menggantikan keputusan UAT.",
    )
    existing.text = (
        "SUS diberikan setelah peserta menyelesaikan tugas relevan. "
        "Instrumen, prosedur penskoran, dan format pelaporan dijelaskan pada Subbab 3.10.3. "
        "Ambang interpretasi awal rerata minimal 68 (Bangor et al., 2008). "
        "Skor SUS tidak menggantikan keputusan UAT."
    )
    for run in existing.runs:
        set_run_font(run)
    set_paragraph_format(existing)

    bab3_target = find_paragraph(doc, "3.11 Pengujian Nonfungsional")
    bab3_items = []
    bab3_items.append(
        add_heading_at_end(
            doc,
            "3.10.3 Pengukuran Usability dengan System Usability Scale (SUS)",
            3,
        )
    )
    bab3_items.append(
        add_body_at_end(
            doc,
            "System Usability Scale (SUS) digunakan untuk mengukur persepsi usability "
            "aplikasi setelah peserta menyelesaikan tugas pada skenario UAT. SUS "
            "menilai persepsi kemudahan, konsistensi, dan keyakinan pengguna; SUS "
            "tidak mengukur kebenaran algoritma maupun keberhasilan fungsi bisnis. "
            "Karena itu, skor SUS dilaporkan terpisah dari keputusan UAT.",
        )
    )
    bab3_items.append(
        add_body_at_end(
            doc,
            "Peserta dikelompokkan menjadi pelanggan dan administrator. Peserta "
            "pelanggan menjalankan login, pencarian/detail produk, pengelolaan "
            "keranjang, checkout, rekomendasi, dan ulasan. Peserta administrator "
            "menjalankan login, pengelolaan produk dan transaksi, upload data, "
            "hitung ulang similarity, serta melihat laporan. Setelah tugas selesai, "
            "peserta mengisi kuesioner secara mandiri; moderator hanya menjelaskan "
            "instruksi dan tidak mengarahkan pilihan jawaban. Kode peserta, peran, "
            "tanggal, perangkat/browser, waktu, error, bantuan, dan jawaban item "
            "dicatat tanpa menampilkan identitas pribadi.",
        )
    )
    bab3_items.append(
        add_body_at_end(
            doc,
            "Setiap pernyataan dijawab dengan skala Likert 1 sampai 5: 1 = sangat "
            "tidak setuju, 2 = tidak setuju, 3 = netral, 4 = setuju, dan 5 = "
            "sangat setuju. Pernyataan bernomor ganjil bersifat positif, sedangkan "
            "pernyataan bernomor genap bersifat negatif.",
        )
    )

    sus_items = [
        (1, "Saya berpikir akan sering menggunakan sistem ini.", "□", "□", "□", "□", "□"),
        (2, "Saya merasa sistem ini terlalu rumit untuk digunakan.", "□", "□", "□", "□", "□"),
        (3, "Saya merasa sistem ini mudah digunakan.", "□", "□", "□", "□", "□"),
        (4, "Saya memerlukan bantuan orang teknis untuk dapat menggunakan sistem ini.", "□", "□", "□", "□", "□"),
        (5, "Saya menemukan berbagai fungsi dalam sistem ini terintegrasi dengan baik.", "□", "□", "□", "□", "□"),
        (6, "Saya merasa terdapat banyak ketidakkonsistenan dalam sistem ini.", "□", "□", "□", "□", "□"),
        (7, "Saya membayangkan kebanyakan orang akan mempelajari cara menggunakan sistem ini dengan sangat cepat.", "□", "□", "□", "□", "□"),
        (8, "Saya merasa sistem ini sangat membebankan untuk digunakan.", "□", "□", "□", "□", "□"),
        (9, "Saya merasa sangat yakin menggunakan sistem ini.", "□", "□", "□", "□", "□"),
        (10, "Saya perlu mempelajari banyak hal terlebih dahulu sebelum menggunakan sistem ini.", "□", "□", "□", "□", "□"),
    ]
    bab3_items.append(add_caption_at_end(doc, "Tabel SUS-1. Kuesioner System Usability Scale"))
    bab3_items.append(
        add_table_at_end(
            doc,
            ["No", "Pernyataan", "1", "2", "3", "4", "5"],
            sus_items,
            widths_cm=[0.7, 9.7, 0.7, 0.7, 0.7, 0.7, 0.7],
        )
    )
    bab3_items.append(add_body_at_end(doc, "Keterangan: beri tanda centang (√) pada satu pilihan jawaban untuk setiap pernyataan."))
    bab3_items.append(add_heading_at_end(doc, "3.10.3.1 Prosedur Perhitungan Skor SUS", 4))
    bab3_items.append(
        add_body_at_end(
            doc,
            "Jika xᵢ adalah jawaban peserta pada item ke-i, kontribusi item ganjil "
            "dihitung dengan cᵢ = xᵢ − 1, sedangkan kontribusi item genap dihitung "
            "dengan cᵢ = 5 − xᵢ. Skor SUS setiap peserta kemudian dihitung dengan "
            "rumus: Skor SUS = (c₁ + c₂ + c₃ + c₄ + c₅ + c₆ + c₇ + c₈ + c₉ + c₁₀) × 2,5. "
            "Skor berada pada rentang 0–100. Rerata keseluruhan dihitung dari seluruh "
            "skor peserta, tanpa membuang skor rendah dan tanpa mengganti jawaban yang kosong.",
        )
    )
    bab3_items.append(
        add_body_at_end(
            doc,
            "Pelaporan SUS mencakup jumlah peserta (n), rerata, median, simpangan baku, "
            "rentang minimum–maksimum, serta rekap deskriptif berdasarkan peran. "
            "Rerata minimal 68 digunakan sebagai ambang interpretasi awal, bukan sebagai "
            "satu-satunya keputusan lulus atau gagal. Nilai di bawah ambang menjadi dasar "
            "untuk menelusuri item bermasalah bersama data penyelesaian tugas, waktu, error, "
            "dan bantuan pengguna (Bangor et al., 2008).",
        )
    )
    move_items_before(bab3_items, bab3_target)

    bab4_target = find_paragraph(doc, "4.8 Batas Bukti dan Pengujian Lanjutan")
    bab4_items = []
    bab4_items.append(add_heading_at_end(doc, "4.7.1 Hasil Pengukuran System Usability Scale (SUS)", 3))
    bab4_items.append(
        add_body_at_end(
            doc,
            "Pada saat dokumen ini disusun belum tersedia berita acara pelaksanaan, "
            "identitas/kode peserta, maupun jawaban 10 item SUS. Oleh sebab itu, skor "
            "numerik SUS belum dapat dihitung dan tidak boleh diisi dengan angka asumsi. "
            "Tabel berikut merupakan format rekap hasil yang diisi setelah pengujian "
            "langsung pada aplikasi.",
        )
    )
    bab4_items.append(add_caption_at_end(doc, "Tabel SUS-2. Rekapitulasi hasil pengukuran SUS"))
    bab4_items.append(
        add_table_at_end(
            doc,
            ["Indikator", "Keseluruhan", "Pelanggan", "Administrator", "Status"],
            [
                ["Jumlah responden (n)", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dilaksanakan"],
                ["Rerata skor SUS", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dapat dihitung"],
                ["Median skor SUS", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dapat dihitung"],
                ["Minimum–maksimum", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dapat dihitung"],
                ["Simpangan baku", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dapat dihitung"],
                ["Penyelesaian tugas usability", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dilaksanakan"],
                ["Waktu median per tugas", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dilaksanakan"],
                ["Jumlah error dan bantuan", "Belum tersedia", "Belum tersedia", "Belum tersedia", "Belum dilaksanakan"],
            ],
            widths_cm=[4.1, 2.4, 2.4, 2.4, 3.0],
        )
    )
    bab4_items.append(
        add_body_at_end(
            doc,
            "Setelah respons tersedia, skor tiap peserta dihitung dengan rumus "
            "[(I1−1) + (5−I2) + (I3−1) + (5−I4) + (I5−1) + (5−I6) + "
            "(I7−1) + (5−I8) + (I9−1) + (5−I10)] × 2,5. Selanjutnya skor peserta "
            "diringkas secara keseluruhan dan berdasarkan peran. Pembahasan harus "
            "menyebutkan jumlah responden dan statistik yang dipakai agar pembaca "
            "dapat menilai ketidakpastian hasil.",
        )
    )
    bab4_items.append(add_heading_at_end(doc, "4.7.2 Interpretasi Hasil SUS", 3))
    bab4_items.append(
        add_body_at_end(
            doc,
            "Jika rerata SUS mencapai minimal 68, usability dipandang memenuhi benchmark "
            "awal untuk pembahasan penelitian ini; jika di bawah 68, antarmuka dan alur "
            "yang memperoleh skor persepsi rendah perlu ditinjau. Interpretasi tersebut "
            "bersifat deskriptif dan tidak menggantikan UAT. Kesimpulan usability juga "
            "harus mempertimbangkan apakah peserta menyelesaikan tugas, berapa lama waktu "
            "yang diperlukan, jumlah error, dan banyaknya bantuan yang diterima.",
        )
    )
    move_items_before(bab4_items, bab4_target)

    # Perbarui status ringkas agar menyebut instrumen telah tersedia, tetapi hasil belum ada.
    status_paragraph = find_paragraph(
        doc,
        "Tabel berikut digunakan untuk membedakan uji yang sudah dibuktikan secara teknis dengan UAT, SUS, kompatibilitas, performa, dan keamanan baseline yang masih memerlukan bukti responden atau eksekusi khusus.",
    )
    status_paragraph.text = (
        "Arsip proyek menyediakan matriks keterlacakan, skenario UAT, dan instrumen SUS "
        "yang kini dilengkapi pada Bab III, tetapi tidak memuat berita acara, record eksekusi, "
        "tangkapan layar, daftar temuan, atau jawaban SUS. Oleh karena itu penerimaan pengguna "
        "dan skor usability belum dapat dinyatakan. Format rekap pada Subbab 4.7.1 hanya boleh "
        "diisi setelah pengujian langsung dilaksanakan."
    )
    for run in status_paragraph.runs:
        set_run_font(run)
    set_paragraph_format(status_paragraph)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
