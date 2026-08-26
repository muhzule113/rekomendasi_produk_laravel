# -*- coding: utf-8 -*-
"""Isi hasil SUS dari ekspor Google Sheets ke dokumen Bab III–IV."""

import statistics
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from openpyxl import load_workbook


ROOT = Path(__file__).resolve().parents[1]
CUSTOMER_XLSX = Path(r"C:\Users\Muh\AppData\Local\Temp\sus_pelanggan.xlsx")
ADMIN_XLSX = Path(r"C:\Users\Muh\AppData\Local\Temp\sus_admin.xlsx")
SOURCE = ROOT / "docs" / "laporan" / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_REVISI_SUS.docx"
)
OUTPUT = ROOT / "docs" / "laporan" / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_REVISI_SUS_HASIL.docx"
)


QUESTION_TEXT = [
    "Saya ingin sering menggunakan sistem ini.",
    "Saya merasa sistem ini terlalu rumit untuk digunakan.",
    "Saya merasa sistem ini mudah digunakan.",
    "Saya membutuhkan bantuan orang lain untuk menggunakan sistem ini.",
    "Fitur-fitur sistem ini terintegrasi dengan baik.",
    "Saya merasa ada terlalu banyak ketidakkonsistenan pada sistem ini.",
    "Saya merasa kebanyakan orang akan cepat mempelajari sistem ini.",
    "Saya merasa sistem ini membingungkan untuk digunakan.",
    "Saya merasa percaya diri menggunakan sistem ini.",
    "Saya perlu belajar banyak sebelum dapat menggunakan sistem ini.",
]


def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def format_paragraph(paragraph, indent=True):
    if indent:
        paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5


def replace_paragraph(paragraph, text, indent=True):
    paragraph.text = text
    format_paragraph(paragraph, indent=indent)
    for run in paragraph.runs:
        set_run_font(run)


def find_paragraph(doc, text):
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise ValueError(f"Paragraf target tidak ditemukan: {text}")


def load_responses(path):
    workbook = load_workbook(path, data_only=True)
    rows = list(workbook.active.values)
    responses = []
    for row in rows[1:]:
        values = [int(value) for value in row[1:11]]
        if len(values) != 10 or any(value < 1 or value > 5 for value in values):
            raise ValueError(f"Respons SUS tidak valid pada {path}: {values}")
        contributions = [
            value - 1 if index % 2 == 0 else 5 - value
            for index, value in enumerate(values)
        ]
        responses.append(
            {
                "values": values,
                "contributions": contributions,
                "score": sum(contributions) * 2.5,
            }
        )
    return responses


def fmt(value):
    return f"{value:.2f}".replace(".", ",")


def stats_for(responses):
    scores = [item["score"] for item in responses]
    return {
        "n": len(scores),
        "mean": statistics.mean(scores),
        "median": statistics.median(scores),
        "minimum": min(scores),
        "maximum": max(scores),
        "stdev": statistics.stdev(scores) if len(scores) > 1 else None,
    }


def set_cell(cell, value, size=9, bold=False, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(str(value))
    set_run_font(run, size=size, bold=bold)


def set_table_header_repeat(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_result_table_before(doc, target, caption, headers, rows, widths):
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.space_after = Pt(6)
    run = caption_paragraph.add_run(caption)
    set_run_font(run, size=11, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    set_table_header_repeat(table.rows[0])
    for index, header in enumerate(headers):
        set_cell(
            table.rows[0].cells[index],
            header,
            size=9,
            bold=True,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
    for row_index, row in enumerate(rows, start=1):
        for column_index, value in enumerate(row):
            alignment = WD_ALIGN_PARAGRAPH.CENTER if column_index != 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell(table.rows[row_index].cells[column_index], value, size=9, align=alignment)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Cm(width)
    for index, width in enumerate(widths):
        table.columns[index].width = Cm(width)

    target._p.addprevious(caption_paragraph._p)
    target._p.addprevious(table._tbl)


def main():
    customer = load_responses(CUSTOMER_XLSX)
    admin = load_responses(ADMIN_XLSX)
    all_responses = customer + admin
    overall_stats = stats_for(all_responses)
    customer_stats = stats_for(customer)
    admin_stats = stats_for(admin)

    document = Document(SOURCE)

    # Samakan instrumen di Bab III dengan pernyataan yang benar-benar dipakai pada Form.
    questionnaire = next(
        table
        for table in document.tables
        if table.cell(0, 0).text == "No"
        and table.cell(0, 1).text == "Pernyataan"
    )
    for index, text in enumerate(QUESTION_TEXT, start=1):
        questionnaire.cell(index, 1).text = text
        for run in questionnaire.cell(index, 1).paragraphs[0].runs:
            set_run_font(run, size=9)

    # Perbarui status ringkas pada Bab IV.
    status_paragraph = find_paragraph(
        document,
        "Arsip proyek menyediakan matriks keterlacakan, skenario UAT, dan instrumen SUS yang kini dilengkapi pada Bab III, tetapi tidak memuat berita acara, record eksekusi, tangkapan layar, daftar temuan, atau jawaban SUS. Oleh karena itu penerimaan pengguna dan skor usability belum dapat dinyatakan. Format rekap pada Subbab 4.7.1 hanya boleh diisi setelah pengujian langsung dilaksanakan.",
    )
    replace_paragraph(
        status_paragraph,
        "Arsip proyek kini memiliki rekap 11 respons SUS dari dua kelompok pengguna, yaitu 10 pelanggan dan 1 administrator. Hasil SUS dapat dihitung dari respons kuesioner, sedangkan UAT, bukti penyelesaian tugas, waktu, error, bantuan, dan keputusan penerimaan pengguna tetap belum tersedia. Skor administrator ditafsirkan secara hati-hati karena hanya terdapat satu responden.",
    )

    table_status = next(
        table for table in document.tables if table.cell(0, 0).text == "Pengujian"
    )
    sus_status_row = next(row for row in table_status.rows if row.cells[0].text == "SUS")
    set_cell(sus_status_row.cells[2], "11 respons: 10 pelanggan dan 1 admin", size=9)
    set_cell(
        sus_status_row.cells[3],
        "Skor per responden dan rekap statistik diisi; interpretasi admin dibatasi karena n=1.",
        size=9,
    )

    summary = next(
        table for table in document.tables if table.cell(0, 0).text == "Indikator"
    )
    summary_rows = {
        "Jumlah responden (n)": ["11", "10", "1", "Selesai dihitung"],
        "Rerata skor SUS": [fmt(overall_stats["mean"]), fmt(customer_stats["mean"]), fmt(admin_stats["mean"]), "Selesai dihitung"],
        "Median skor SUS": [fmt(overall_stats["median"]), fmt(customer_stats["median"]), fmt(admin_stats["median"]), "Selesai dihitung"],
        "Minimum–maksimum": [
            f"{fmt(overall_stats['minimum'])}–{fmt(overall_stats['maximum'])}",
            f"{fmt(customer_stats['minimum'])}–{fmt(customer_stats['maximum'])}",
            f"{fmt(admin_stats['minimum'])}–{fmt(admin_stats['maximum'])}",
            "Selesai dihitung",
        ],
        "Simpangan baku": [fmt(overall_stats["stdev"]), fmt(customer_stats["stdev"]), "– (n=1)", "Sampel; admin deskriptif"],
        "Penyelesaian tugas usability": ["Belum tersedia", "Belum tersedia", "Belum tersedia", "Tidak tercatat pada spreadsheet"],
        "Waktu median per tugas": ["Belum tersedia", "Belum tersedia", "Belum tersedia", "Tidak tercatat pada spreadsheet"],
        "Jumlah error dan bantuan": ["Belum tersedia", "Belum tersedia", "Belum tersedia", "Tidak tercatat pada spreadsheet"],
    }
    for row in summary.rows[1:]:
        label = row.cells[0].text
        values = summary_rows[label]
        for index, value in enumerate(values, start=1):
            set_cell(row.cells[index], value, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Masukkan tabel skor anonim per responden sebagai bukti perhitungan.
    score_rows = []
    for prefix, responses in (("P", customer), ("A", admin)):
        for index, response in enumerate(responses, start=1):
            score = response["score"]
            note = "Memenuhi benchmark ≥ 68" if score >= 68 else "Di bawah benchmark 68"
            if prefix == "A":
                note += "; interpretasi terbatas n=1"
            score_rows.append(
                [
                    f"{prefix}-{index:02d}",
                    "Pelanggan" if prefix == "P" else "Admin",
                    ", ".join(str(value) for value in response["values"]),
                    fmt(score),
                    note,
                ]
            )
    formula_paragraph = find_paragraph(
        document,
        "Setelah respons tersedia, skor tiap peserta dihitung dengan rumus [(I1−1) + (5−I2) + (I3−1) + (5−I4) + (I5−1) + (5−I6) + (I7−1) + (5−I8) + (I9−1) + (5−I10)] × 2,5. Selanjutnya skor peserta diringkas secara keseluruhan dan berdasarkan peran. Pembahasan harus menyebutkan jumlah responden dan statistik yang dipakai agar pembaca dapat menilai ketidakpastian hasil.",
    )
    add_result_table_before(
        document,
        formula_paragraph,
        "Tabel SUS-3. Skor SUS per responden (kode anonim)",
        ["Kode", "Peran", "Respons I1–I10", "Skor SUS", "Interpretasi awal"],
        score_rows,
        [1.3, 2.0, 5.7, 1.5, 3.5],
    )
    replace_paragraph(
        formula_paragraph,
        "Berdasarkan 11 respons yang tersedia, skor tiap peserta dihitung dengan rumus [(I1−1) + (5−I2) + (I3−1) + (5−I4) + (I5−1) + (5−I6) + (I7−1) + (5−I8) + (I9−1) + (5−I10)] × 2,5. Rerata keseluruhan adalah 77,50; rerata pelanggan 80,25; dan rerata administrator 50,00. Nilai administrator tidak dapat digeneralisasi karena hanya berasal dari satu responden.",
    )

    interpretation = find_paragraph(
        document,
        "Jika rerata SUS mencapai minimal 68, usability dipandang memenuhi benchmark awal untuk pembahasan penelitian ini; jika di bawah 68, antarmuka dan alur yang memperoleh skor persepsi rendah perlu ditinjau. Interpretasi tersebut bersifat deskriptif dan tidak menggantikan UAT. Kesimpulan usability juga harus mempertimbangkan apakah peserta menyelesaikan tugas, berapa lama waktu yang diperlukan, jumlah error, dan banyaknya bantuan yang diterima.",
    )
    replace_paragraph(
        interpretation,
        "Rerata SUS keseluruhan sebesar 77,50 berada di atas benchmark awal 68 sehingga persepsi usability dari 11 respons tergolong positif pada tingkat agregat. Kelompok pelanggan memperoleh rerata 80,25 dengan median 77,50, sedangkan admin memperoleh skor 50,00. Hasil admin belum dapat dijadikan kesimpulan kelompok karena n=1. Satu responden pelanggan dan satu responden admin berada di bawah benchmark, sehingga alur penggunaan dan kebutuhan bantuan pada kedua pengalaman tersebut perlu ditelusuri pada UAT lanjutan. Interpretasi ini tetap bersifat deskriptif; efektivitas, efisiensi, error, dan bantuan belum dapat disimpulkan karena tidak dicatat pada spreadsheet SUS.",
    )

    boundary = next(
        table
        for table in document.tables
        if any(row.cells[0].text == "UAT + SUS" for row in table.rows)
    )
    sus_boundary_row = next(row for row in boundary.rows if row.cells[0].text == "UAT + SUS")
    set_cell(sus_boundary_row.cells[1], "SUS selesai; UAT belum dilaksanakan", size=9)
    set_cell(
        sus_boundary_row.cells[2],
        "Pertahankan rekap 11 respons SUS; laksanakan UAT dan lampirkan bukti tugas, temuan, serta berita acara.",
        size=9,
    )

    conclusion = find_paragraph(
        document,
        "4. Skenario black-box, UAT, SUS, load test endpoint, dan keamanan baseline telah dirancang dari implementasi aktual tetapi belum dilaksanakan pada lingkungan penyusunan dokumen; sistem belum dapat dinyatakan diterima pengguna atau aman/performan hanya dari inspeksi kode.",
    )
    replace_paragraph(
        conclusion,
        "4. SUS telah diisi oleh 11 responden dan menghasilkan rerata 77,50; pelanggan memperoleh rerata 80,25, sedangkan admin 50,00 dari satu responden sehingga tidak dapat digeneralisasi. UAT, load test endpoint, dan keamanan baseline masih memerlukan pelaksanaan serta bukti langsung; sistem belum dapat dinyatakan diterima pengguna hanya dari skor SUS.",
    )

    document.save(OUTPUT)
    print(OUTPUT)
    print(
        f"overall={fmt(overall_stats['mean'])}; customer={fmt(customer_stats['mean'])}; admin={fmt(admin_stats['mean'])}"
    )


if __name__ == "__main__":
    main()
