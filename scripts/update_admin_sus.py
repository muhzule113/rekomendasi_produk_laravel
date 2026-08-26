# -*- coding: utf-8 -*-
"""Perbarui respons administrator dan rekap SUS pada DOCX hasil."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "laporan" / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_REVISI_SUS_HASIL.docx"
)
OUTPUT = ROOT / "docs" / "laporan" / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_REVISI_SUS_HASIL_ADMIN_UPDATED.docx"
)


def set_run_font(run, size=9, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def set_cell(cell, text, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=size)


def find_paragraph_containing(document, fragment):
    for paragraph in document.paragraphs:
        if fragment in paragraph.text:
            return paragraph
    raise ValueError(f"Paragraf tidak ditemukan: {fragment}")


def replace_paragraph(paragraph, text):
    paragraph.text = text
    paragraph.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run, size=11)


def main():
    document = Document(SOURCE)

    summary = next(
        table for table in document.tables if table.cell(0, 0).text == "Indikator"
    )
    replacement = {
        "Jumlah responden (n)": ["11", "10", "1", "Selesai dihitung"],
        "Rerata skor SUS": ["80,91", "80,25", "87,50", "Selesai dihitung"],
        "Median skor SUS": ["80,00", "77,50", "87,50", "Selesai dihitung"],
        "Minimum–maksimum": ["50,00–100,00", "50,00–100,00", "87,50–87,50", "Selesai dihitung"],
        "Simpangan baku": ["15,05", "15,70", "– (n=1)", "Sampel; admin deskriptif"],
    }
    for row in summary.rows[1:]:
        label = row.cells[0].text
        if label in replacement:
            for index, value in enumerate(replacement[label], start=1):
                set_cell(row.cells[index], value)

    respondent_table = next(
        table for table in document.tables if table.cell(0, 0).text == "Kode"
    )
    admin_row = next(row for row in respondent_table.rows if row.cells[0].text == "A-01")
    set_cell(admin_row.cells[2], "5, 2, 5, 2, 5, 2, 5, 2, 5, 2")
    set_cell(admin_row.cells[3], "87,50")
    set_cell(admin_row.cells[4], "Memenuhi benchmark ≥ 68; interpretasi terbatas n=1")

    formula = find_paragraph_containing(document, "Berdasarkan 11 respons yang tersedia")
    replace_paragraph(
        formula,
        "Berdasarkan 11 respons yang tersedia, skor tiap peserta dihitung dengan rumus [(I1−1) + (5−I2) + (I3−1) + (5−I4) + (I5−1) + (5−I6) + (I7−1) + (5−I8) + (I9−1) + (5−I10)] × 2,5. Rerata keseluruhan adalah 80,91 dengan median 80,00; rerata pelanggan adalah 80,25; dan rerata administrator adalah 87,50. Nilai administrator tidak dapat digeneralisasi karena hanya berasal dari satu responden.",
    )

    interpretation = find_paragraph_containing(
        document, "Rerata SUS keseluruhan sebesar 77,50"
    )
    replace_paragraph(
        interpretation,
        "Rerata SUS keseluruhan sebesar 80,91 berada di atas benchmark awal 68 sehingga persepsi usability dari 11 respons tergolong positif pada tingkat agregat. Kelompok pelanggan memperoleh rerata 80,25 dengan median 77,50, sedangkan admin memperoleh skor 87,50. Hasil admin belum dapat dijadikan kesimpulan kelompok karena n=1. Satu responden pelanggan berada di bawah benchmark, sehingga alur penggunaan dan kebutuhan bantuan pada pengalaman tersebut perlu ditelusuri pada UAT lanjutan. Interpretasi ini tetap bersifat deskriptif; efektivitas, efisiensi, error, dan bantuan belum dapat disimpulkan karena tidak dicatat pada spreadsheet SUS.",
    )

    conclusion = find_paragraph_containing(
        document, "4. SUS telah diisi oleh 11 responden"
    )
    replace_paragraph(
        conclusion,
        "4. SUS telah diisi oleh 11 responden dan menghasilkan rerata 80,91; pelanggan memperoleh rerata 80,25, sedangkan admin 87,50 dari satu responden sehingga tidak dapat digeneralisasi. UAT, load test endpoint, dan keamanan baseline masih memerlukan pelaksanaan serta bukti langsung; sistem belum dapat dinyatakan diterima pengguna hanya dari skor SUS.",
    )

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
