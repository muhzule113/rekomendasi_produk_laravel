# -*- coding: utf-8 -*-
"""Generate BAB III & IV pengujian yang selaras dengan kode repositori."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT = Path(__file__).resolve().parents[1] / (
    "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap.docx"
)


def set_run_font(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=12 if level == 1 else 11, bold=True)
    return p


def add_para(doc, text, first_line_indent=True):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run)
    return p


def add_para_plain(doc, text):
    return add_para(doc, text, first_line_indent=False)


def add_table(doc, headers, rows, caption=None):
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=11, bold=True)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
    doc.add_paragraph()
    return table


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(3)
    section.left_margin = Cm(4)
    section.right_margin = Cm(3)

    # ========== BAB III ==========
    add_heading_custom(doc, "BAB III. METODE PENGUJIAN", 1)

    add_heading_custom(doc, "3.1 Jenis dan Tujuan Pengujian", 2)
    add_para(
        doc,
        "Pengujian sistem dalam penelitian ini menilai dua lapisan: "
        "(1) evaluasi analitik kuantitatif secara offline terhadap kualitas "
        "rekomendasi Item-Based Collaborative Filtering (IBCF), dan "
        "(2) pengujian perangkat lunak berbasis repositori yang mencakup "
        "tes otomatis, skenario black-box/integrasi, serta rencana User "
        "Acceptance Testing (UAT) dan usability. Evaluasi offline dipilih "
        "karena dapat diulang pada snapshot data yang sama tanpa "
        "memperlihatkan rekomendasi eksperimental kepada pelanggan "
        "(Herlocker et al., 2004; Zangerle & Bauer, 2022).",
    )
    add_para(
        doc,
        "Metode utama yang diuji adalah IBCF dengan implicit feedback biner "
        "dan metrik Cosine Similarity, sesuai implementasi pada "
        "RecommenderService (PHP) dan python/cf/ (Python). Tujuan pengujian "
        "adalah memverifikasi kebenaran matematis cosine, ketiadaan data "
        "leakage pada holdout temporal, perilaku cold-start/fallback, "
        "kelulusan suite tes otomatis, serta kesiapan skenario UAT—bukan "
        "mengklaim penerimaan pengguna sebelum skenario tersebut dijalankan.",
    )

    add_heading_custom(doc, "3.2 Sumber Data, Unit Analisis, dan Privasi", 2)
    add_para(
        doc,
        "Sumber utama evaluasi offline yang selaras dengan aplikasi adalah "
        "snapshot basis data transaksi yang memenuhi filter produksi: "
        "status_pesanan = Selesai dan status_pembayaran = Dibayar, serta "
        "produk berstatus aktif. Data dimuat melalui "
        "python/cf/data_loader.py (load_interaction_events). Unit analisis "
        "adalah pelanggan (id_user) setelah resolusi identitas pada pipeline "
        "unggah (email/telepon dinormalisasi).",
    )
    add_para(
        doc,
        "Apabila digunakan berkas CSV eksternal untuk eksperimen, kolom yang "
        "tersedia dan filter yang dapat diterapkan harus dicatat secara "
        "eksplisit. Jika CSV tidak memuat status_pembayaran, filter Dibayar "
        "tidak dapat diterapkan pada berkas tersebut; perbedaan ini merupakan "
        "keterbatasan reproduksi terhadap konfigurasi produksi dan wajib "
        "dilaporkan. Email dan nomor telepon diperlakukan sebagai data "
        "pribadi: hanya dipakai untuk pengelompokan internal dan tidak "
        "ditampilkan pada contoh hasil (diganti pseudonim bila diperlukan).",
    )

    add_heading_custom(doc, "3.3 Seleksi dan Pembersihan Data", 2)
    add_para(
        doc,
        "Baris yang dipakai membangun model pada jalur produksi dan evaluator "
        "berbasis basis data harus memenuhi: transaksi Selesai + Dibayar, "
        "produk aktif, serta kuantitas/harga sah sesuai aturan pipeline ETL. "
        "Interaksi pelanggan–produk dibentuk secara biner. Pembelian ulang "
        "dan jumlah barang tidak menambah bobot preferensi; setiap pasangan "
        "pelanggan–produk hanya membentuk satu nilai interaksi "
        "(x_ui = 1 jika pernah membeli pada transaksi valid; selain itu 0). "
        "Pendekatan ini sesuai IBCF berbasis implicit feedback dan mencegah "
        "satu pelanggan mendominasi perhitungan kemiripan.",
    )

    add_heading_custom(
        doc, "3.4 Pemisahan Data Temporal (Leave-Last-Item Holdout)", 2
    )
    add_para(
        doc,
        "Pengujian offline mengikuti implementasi python/cf/evaluator.py "
        "(prepare_holdout_splits). Untuk setiap pelanggan, interaksi diurutkan "
        "berdasarkan tanggal, kemudian id_transaction, kemudian id_product. "
        "Urutan kemunculan pertama setiap produk dicatat. Produk dengan "
        "kemunculan pertama paling akhir dijadikan satu item uji (holdout); "
        "produk sebelumnya menjadi data latih. Pelanggan dievaluasi hanya "
        "jika memiliki sedikitnya dua produk berbeda dan setidaknya satu "
        "produk latih setelah holdout dikeluarkan.",
    )
    add_para(
        doc,
        "Matriks interaksi, co-occurrence, dan cosine similarity dibangun "
        "hanya dari interaksi latih (ditambah riwayat pelanggan yang tidak "
        "dievaluasi karena kurang dari dua produk). Item holdout tidak ikut "
        "membentuk model sehingga kebocoran data dihindari. Protokol ini "
        "adalah leave-last-item berbasis waktu, bukan leave-last-date-out "
        "(seluruh produk baru pada tanggal terakhir). Karena target uji "
        "berjumlah satu produk per pelanggan, Recall@K per pelanggan bernilai "
        "0 atau 1.",
    )

    add_heading_custom(doc, "3.5 Perhitungan Item-Based Collaborative Filtering", 2)
    add_para(
        doc,
        "Kemiripan dua produk dihitung menggunakan cosine similarity atas "
        "vektor pelanggan biner (Sarwar et al., 2001):",
    )
    add_para_plain(
        doc,
        "sim(i,j) = n(i,j) / sqrt(n(i) × n(j))",
    )
    add_para(
        doc,
        "dengan n(i,j) jumlah pelanggan yang membeli produk i dan j, serta "
        "n(i) dan n(j) jumlah pelanggan masing-masing produk. Pasangan hanya "
        "disimpan jika co-occurrence memenuhi ambang minimum "
        "CF_MIN_CO_OCCURRENCE (default 2). Diagonal matriks dibuat nol dan "
        "hasil harus simetris. Skor tidak dinormalisasi terhadap skor "
        "maksimum katalog.",
    )
    add_para(
        doc,
        "Skor produk kandidat j untuk pelanggan u dihitung sebagai rata-rata "
        "kemiripan terhadap seluruh produk dalam riwayat latih H_u "
        "(sama pada RecommenderService::recommendForCustomer dan "
        "recommend_top_k):",
    )
    add_para_plain(
        doc,
        "score(u,j) = (1 / |H_u|) × Σ_{i∈H_u} sim(i,j)",
    )
    add_para(
        doc,
        "Kandidat harus belum pernah dibeli dan memiliki skor positif. "
        "Peringkat ditentukan oleh: (1) skor prediksi tertinggi, "
        "(2) co-occurrence / jumlah sisi pendukung tertinggi, "
        "(3) id_product secara menaik. Aturan tersebut membuat hasil "
        "deterministik dan selaras antara PHP dan Python.",
    )

    add_heading_custom(doc, "3.6 Metrik Evaluasi Offline", 2)
    add_para(
        doc,
        "Modul evaluasi produksi (python/cf/evaluator.py, fungsi "
        "evaluate_at_k dan run_evaluation) menghitung metrik berikut untuk "
        "K = 5 dan K = 10, lalu dapat menyimpan hasil ke tabel "
        "evaluation_logs dengan method ibcf_cosine_time_holdout.",
    )
    add_table(
        doc,
        ["Metrik", "Perhitungan", "Makna"],
        [
            [
                "Precision@K",
                "hit / K (hit = 1 jika item holdout masuk Top-K)",
                "Ketepatan isi daftar rekomendasi",
            ],
            [
                "Recall@K",
                "hit / 1",
                "Keberhasilan menemukan satu item holdout",
            ],
            [
                "F1@K",
                "2 × P × R / (P + R)",
                "Keseimbangan precision dan recall",
            ],
            [
                "Hit Rate@K",
                "Proporsi pelanggan dengan hit = 1",
                "Keberhasilan minimal per pelanggan",
            ],
            [
                "Catalog Coverage@K",
                "Produk unik yang direkomendasikan / ukuran katalog × 100%",
                "Luas jangkauan produk",
            ],
        ],
        caption="Tabel 1. Definisi metrik pengujian offline pada modul produksi",
    )
    add_para(
        doc,
        "Metrik dihitung per pelanggan lalu dirata-ratakan secara makro. "
        "Pelanggan tanpa rekomendasi tetap masuk denominator dengan nilai "
        "nol bila Top-K kosong. NDCG, baseline Most Popular, acuan peluang "
        "acak, dan uji paired bootstrap tidak diimplementasikan pada "
        "python/cf/evaluator.py; metrik tersebut tidak dilaporkan sebagai "
        "keluaran modul produksi kecuali skrip terpisah dilampirkan dan "
        "dijalankan secara eksplisit.",
    )

    add_heading_custom(doc, "3.7 Konfigurasi Reproduksibilitas", 2)
    add_table(
        doc,
        ["Parameter", "Nilai"],
        [
            ["Modul evaluasi", "python/cf/evaluator.py"],
            ["Filter interaksi", "status_pesanan = Selesai AND status_pembayaran = Dibayar; produk aktif"],
            ["Pengenal pelanggan", "id_user (setelah resolusi pipeline)"],
            ["Interaksi", "Biner (pernah membeli pada transaksi valid)"],
            ["Pemisahan", "Leave-last-item per pelanggan; 1 produk holdout"],
            ["Similarity", "Cosine biner; CF_MIN_CO_OCCURRENCE = 2"],
            ["Agregasi skor", "sum(sim) / |H_u|"],
            ["Tie-break", "prediction_score ↓, co_occurrence ↓, id_product ↑"],
            ["K", "5 dan 10"],
            ["Label method", "ibcf_cosine_time_holdout"],
            ["Persistensi", "Tabel evaluation_logs (opsional)"],
        ],
        caption="Tabel 2. Konfigurasi reproduksibilitas evaluasi offline",
    )

    add_heading_custom(doc, "3.8 Pengujian Perangkat Lunak Berbasis Repositori", 2)
    add_para(
        doc,
        "Pengujian perangkat lunak diturunkan dari implementasi aktual pada "
        "arsip proyek. Inventaris mencakup tes unit PHP, tes fitur Laravel, "
        "serta tes unit Python untuk cosine similarity, evaluator, parity "
        "fixture, dan pipeline. Pengujian otomatis digunakan sebagai bukti "
        "regresi; pengujian black-box memeriksa alur antarmuka, otorisasi, "
        "integrasi basis data, pembayaran, dan unggah data. Keberadaan file "
        "tes tidak dianggap bukti lulus sebelum suite dieksekusi pada runtime "
        "yang memenuhi dependensi proyek.",
    )

    add_heading_custom(doc, "3.8.1 Pengujian Otomatis", 3)
    add_table(
        doc,
        ["Lapisan", "Berkas utama", "Fokus pemeriksaan", "Bukti kelulusan"],
        [
            [
                "Unit PHP",
                "CosineSimilarityTest",
                "Fixture cosine, vektor identik, irisan nol, simetri, rentang skor, tanpa diagonal",
                "Assertion numerik lulus (toleransi 10⁻⁶)",
            ],
            [
                "Feature Laravel",
                "RecommenderServiceTest",
                "Matriks transaksi valid, min co-occurrence, agregasi kandidat, cold start, otorisasi admin, dirty flag",
                "Respons HTTP / state DB / skor sesuai fixture",
            ],
            [
                "Feature Laravel",
                "ReviewTest",
                "Satu ulasan per pelanggan–produk; penolakan duplikasi",
                "Ulasan pertama tersimpan; duplikat ditolak",
            ],
            [
                "Unit Python",
                "test_cosine.py",
                "Cosine biner, tanpa diagonal, threshold, simetri, tanpa normalisasi maksimum",
                "Seluruh unittest lulus",
            ],
            [
                "Evaluator Python",
                "test_evaluator.py",
                "Leave-last-item berbasis waktu, Top-K tanpa duplikasi, pencegahan leakage holdout",
                "Split, rekomendasi, dan metrik sesuai fixture",
            ],
            [
                "Parity Python",
                "test_parity_fixture.py",
                "Kesetaraan vektor dan aturan penyimpanan pasangan terhadap expected PHP",
                "Selisih skor dalam toleransi",
            ],
            [
                "Pipeline Python",
                "test_pipeline.py",
                "Penolakan user tidak dikenal, grouping kode transaksi, deteksi kolom wajib",
                "Record invalid ditolak; transformasi sesuai aturan",
            ],
        ],
        caption="Tabel 3. Lapisan pengujian otomatis pada repositori",
    )
    add_para(
        doc,
        "Inventaris repositori memuat 29 metode tes otomatis: 13 tes Python "
        "dan 16 tes PHP/Laravel (termasuk dua berkas contoh bawaan Laravel "
        "dan satu assertion placeholder yang perlu dibedakan dari tes domain "
        "substantif). Suite dijalankan pada basis data pengujian terisolasi; "
        "kelulusan ditentukan oleh assertion relevan, ketiadaan exception "
        "tak terduga, dan exit code nol.",
    )

    add_heading_custom(doc, "3.8.2 Pengujian Black-Box dan Integrasi", 3)
    add_para(
        doc,
        "Skenario black-box disusun dari route web/API, middleware role, "
        "controller, dan service. Setiap kasus harus dijalankan melalui "
        "antarmuka atau request HTTP terhadap aplikasi uji; pembacaan source "
        "code hanya menjadi dasar penyusunan skenario, bukan hasil pengujian.",
    )
    add_table(
        doc,
        ["ID", "Modul", "Skenario inti", "Hasil yang diharapkan"],
        [
            ["BB-01", "Registrasi", "Data valid, email duplikat, konfirmasi berbeda, password < 6", "Valid tersimpan; invalid ditolak dengan pesan tepat"],
            ["BB-02", "Login dan role", "Login pelanggan/admin, role salah, password salah, akun nonaktif", "Diarahkan sesuai role; akses tidak sah ditolak"],
            ["BB-03", "Keranjang", "Tambah, ubah jumlah, hapus, login untuk pulihkan sesi", "Isi keranjang konsisten dan tersinkron"],
            ["BB-04", "Checkout", "Keranjang valid, kosong, jumlah > stok", "Transaksi valid tercatat; invalid tidak mengurangi stok"],
            ["BB-05", "Rekomendasi personal", "Pelanggan dengan riwayat Selesai+Dibayar membuka /rekomendasi", "Produk dibeli tidak muncul; kandidat aktif + alasan IBCF"],
            ["BB-06", "Cold start", "Guest atau pelanggan tanpa riwayat valid", "Fallback berlabel eksplisit bukan collaborative filtering"],
            ["BB-07", "Produk serupa", "Detail produk dengan/ tanpa tetangga similarity", "Produk aktif tersedia; tidak menduplikasi produk asal"],
            ["BB-08", "Ulasan", "Kirim ulasan lalu ulangi pada produk sama", "Pertama diterima; duplikasi ditolak"],
            ["BB-09", "Otorisasi admin", "Akses /admin dan hitung ulang sebagai guest/pelanggan/admin", "Hanya admin terautentikasi yang dapat menjalankan fungsi admin"],
            ["BB-10", "Kelola produk", "Tambah/edit dengan gambar valid/invalid/ukuran berlebih", "Valid disimpan; invalid ditolak; data lama konsisten"],
            ["BB-11", "Upload data", "CSV/XLSX valid, format terlarang, terlalu besar, invalid, hash duplikat", "Valid diproses; invalid/duplikat ditolak dan tercatat di log"],
            ["BB-12", "Similarity", "Hitung ulang pada matriks berisi dan kosong", "Penyimpanan atomik pada matriks valid; kosong tidak menghapus data lama"],
            ["BB-13", "Status transaksi", "Transisi status valid dan yang dilarang", "Hanya transisi valid diterapkan dan tercatat"],
            ["BB-14", "Midtrans", "Checkout + notifikasi/verifikasi sandbox", "Status berubah hanya untuk notifikasi terverifikasi dan order sesuai"],
        ],
        caption="Tabel 4. Skenario pengujian black-box dan integrasi",
    )

    add_heading_custom(doc, "3.9 Pengujian Penerimaan Pengguna dan Usability", 2)
    add_para(
        doc,
        "UAT memverifikasi proses bisnis pelanggan dan administrator sesuai "
        "kebutuhan, sedangkan System Usability Scale (SUS) mengukur persepsi "
        "usability pada skala 0–100 setelah tugas selesai (Brooke, 1996). "
        "Keduanya dilaporkan terpisah agar penerimaan fungsional tidak "
        "dicampur dengan persepsi kemudahan (ISO 9241-11:2018; "
        "ISO/IEC/IEEE 29119-3:2021).",
    )
    add_para(
        doc,
        "UAT dilaksanakan pada lingkungan uji terisolasi dengan basis data, "
        "akun, stok, berkas unggahan, dan kredensial Midtrans sandbox yang "
        "terkontrol. Hasil tiap skenario dicatat sebagai Lulus, Gagal, atau "
        "Tertahan, dilengkapi bukti dan kode temuan. Kriteria keluar antara "
        "lain: seluruh skenario prioritas kritis dieksekusi tanpa temuan "
        "Critical/High terbuka; kelulusan keseluruhan minimal 90%; pemilik "
        "sistem menandatangani keputusan akhir.",
    )

    add_heading_custom(doc, "3.9.1 Matriks Keterlacakan Kebutuhan", 3)
    add_table(
        doc,
        ["ID", "Aktor", "Kebutuhan", "Bukti implementasi", "Kasus UAT"],
        [
            ["RF-01", "Pelanggan", "Registrasi, login, logout", "AuthController; /register, /login, /logout", "P-01–P-03"],
            ["RF-02", "Pelanggan", "Katalog dan detail produk", "ProdukController; /produk, /produk/{id}", "P-04–P-05"],
            ["RF-03", "Pelanggan", "Keranjang belanja", "CartController / CartService", "P-06–P-08"],
            ["RF-04", "Pelanggan", "Checkout bayar di toko", "CheckoutController", "P-09–P-10"],
            ["RF-05", "Pelanggan", "Pembayaran Midtrans", "MidtransPaymentService; verify-payment; webhook", "P-11–P-12"],
            ["RF-06", "Pelanggan", "Riwayat transaksi", "TransaksiController pelanggan", "P-13"],
            ["RF-07", "Pelanggan", "Rekomendasi personal & cold start", "RekomendasiController; RecommenderService", "P-14–P-16"],
            ["RF-08", "Pelanggan", "Produk serupa & ulasan", "Detail produk; ReviewController", "P-17–P-18"],
            ["RF-09", "Admin", "Login & otorisasi admin", "AdminMiddleware; AuthController", "A-01–A-02"],
            ["RF-10", "Admin", "Kelola produk & gambar", "Admin ProdukController; storage", "A-03–A-05"],
            ["RF-11", "Admin", "Kelola transaksi & status", "Admin TransaksiController", "A-06–A-07"],
            ["RF-12", "Admin", "Upload transaksi & pipeline", "UploadService; python/pipeline", "A-08–A-10"],
            ["RF-13", "Admin", "Hitung ulang similarity & analisis", "AnalisisController; RecommenderService / CF Python", "A-11–A-13"],
            ["RF-14", "Admin", "Laporan & ulasan", "Laporan / Review admin", "A-14"],
        ],
        caption="Tabel 5. Matriks keterlacakan kebutuhan terhadap skenario UAT",
    )

    add_heading_custom(doc, "3.9.2 Ringkasan Skenario UAT", 3)
    add_para(
        doc,
        "Skenario pelanggan (contoh prioritas kritis/tinggi): registrasi valid "
        "dan invalid; login sesuai role; jelajah katalog; kelola keranjang; "
        "checkout bayar di toko; checkout Midtrans sandbox hingga token/Snap; "
        "riwayat dan sinkronisasi status; rekomendasi berriwayat; cold start "
        "dengan label fallback; produk serupa; satu ulasan tanpa duplikasi.",
    )
    add_para(
        doc,
        "Skenario administrator: penolakan akses non-admin; CRUD produk dengan "
        "validasi gambar; perubahan status transaksi yang sah; unggah "
        "CSV/XLSX valid dan penolakan invalid/duplikat; hitung ulang "
        "similarity pada matriks valid dan kosong; dashboard analisis "
        "menampilkan statistik/log; pengelolaan ulasan dan laporan.",
    )
    add_para(
        doc,
        "SUS diberikan setelah peserta menyelesaikan tugas relevan (10 "
        "pernyataan, skor 0–100). Ambang interpretasi awal rerata minimal 68 "
        "(Bangor et al., 2008). Skor SUS tidak menggantikan keputusan UAT.",
    )

    add_heading_custom(doc, "3.10 Pengujian Nonfungsional", 2)
    add_table(
        doc,
        ["Aspek", "Kriteria / ukuran", "Catatan"],
        [
            ["Otorisasi", "Guest/pelanggan tidak mengakses fungsi admin", "Baseline keamanan berbasis risiko"],
            ["Integritas data", "Checkout/upload gagal tidak meninggalkan record parsial", "Diverifikasi lewat BB dan tes fitur"],
            ["Validasi input", "Input di luar kebijakan ditolak", "Termasuk file upload dan form auth"],
            ["Kinerja endpoint", "Latency p50/p95, throughput, error rate", "Benchmark evaluator Python ≠ load test Laravel"],
            ["Kompatibilitas", "Desktop dan mobile pada browser yang ditetapkan", "Perlu eksekusi dinamis"],
            ["Pemulihan kegagalan", "Pipeline/CF gagal tercatat di log; data lama aman", "Dirty flag / status log"],
        ],
        caption="Tabel 6. Kriteria pengujian nonfungsional",
    )
    add_para(
        doc,
        "Pemeriksaan keamanan pada penelitian ini merupakan baseline berbasis "
        "risiko dan tidak boleh disebut penetration test penuh tanpa "
        "pelaksanaan oleh penguji yang kompeten. Nilai latency, throughput, "
        "dan temuan keamanan hanya boleh ditulis setelah aplikasi aktif diuji.",
    )

    # ========== BAB IV ==========
    add_heading_custom(doc, "BAB IV. HASIL PENGUJIAN DAN PEMBAHASAN", 1)

    add_heading_custom(doc, "4.1 Ruang Lingkup Hasil yang Dilaporkan", 2)
    add_para(
        doc,
        "Bab ini memisahkan (a) hasil yang dapat diverifikasi dari kode dan "
        "tes di repositori, (b) hasil yang memerlukan eksekusi ulang evaluator "
        "pada snapshot basis data, dan (c) pengujian yang belum dilaksanakan. "
        "Pemisahan ini mencegah klaim akurasi atau penerimaan pengguna yang "
        "tidak didukung bukti runtime.",
    )

    add_heading_custom(doc, "4.2 Validasi Matematis Cosine Similarity", 2)
    add_para(
        doc,
        "Kebenaran rumus cosine biner diverifikasi melalui suite unit pada "
        "repositori. CosineSimilarityTest (PHP) dan test_cosine.py (Python) "
        "memeriksa fixture vektor, kasus vektor identik (skor 1), irisan nol "
        "(skor 0), simetri, rentang [0,1], ketiadaan pasangan diagonal, "
        "penerapan minimum co-occurrence, serta larangan normalisasi terhadap "
        "skor maksimum katalog. test_parity_fixture.py membandingkan keluaran "
        "Python terhadap expected yang selaras aturan penyimpanan PHP.",
    )
    add_table(
        doc,
        ["Pemeriksaan", "Sumber bukti", "Status pada penyusunan dokumen"],
        [
            ["Fixture cosine & simetri", "CosineSimilarityTest; test_cosine.py", "Tersedia di repo; perlu dieksekusi di lingkungan lengkap"],
            ["Tanpa diagonal / skor > 0", "RecommenderServiceTest; test_cosine.py", "Tersedia di repo"],
            ["Min co-occurrence = 2", "RecommenderServiceTest; test_cosine.py", "Tersedia di repo"],
            ["Parity PHP–Python (fixture kecil)", "test_parity_fixture.py + expected_parity.json", "Tersedia di repo"],
            ["Parity full-snapshot produksi", "Belum ada skrip/hasil di arsip", "Belum diuji"],
        ],
        caption="Tabel 7. Status validasi matematis dan parity",
    )
    add_para(
        doc,
        "Contoh perhitungan manual untuk pasangan dengan n(i)=4, n(j)=2, "
        "n(i,j)=2: sim = 2 / sqrt(4×2) = 0,707107 (dalam toleransi 10⁻⁶ "
        "terhadap implementasi). Contoh ini selaras fixture pada "
        "RecommenderServiceTest.",
    )

    add_heading_custom(
        doc, "4.3 Hasil Evaluasi Offline (Siap Diisi dari Evaluator Produksi)", 2
    )
    add_para(
        doc,
        "Tabel berikut mengikuti keluaran run_evaluation() pada "
        "python/cf/evaluator.py. Nilai diisi setelah evaluator dijalankan "
        "pada snapshot basis data penelitian yang memenuhi filter Selesai + "
        "Dibayar. Versi dokumen sebelumnya yang memuat NDCG, Most Popular, "
        "interval bootstrap, dan acuan peluang acak tidak lagi digunakan "
        "sebagai hasil modul produksi karena metrik/baseline tersebut tidak "
        "terdapat pada evaluator repositori.",
    )
    add_table(
        doc,
        [
            "Metode",
            "K",
            "Precision",
            "Recall",
            "F1",
            "Hit Rate",
            "Catalog Coverage",
            "Users",
        ],
        [
            ["IBCF Cosine", "5", "…", "…", "…", "…", "…", "…"],
            ["IBCF Cosine", "10", "…", "…", "…", "…", "…", "…"],
        ],
        caption="Tabel 8. Hasil evaluasi kualitas peringkat (isi dari evaluation_logs / output evaluator)",
    )
    add_para(
        doc,
        "Cara mengisi: jalankan evaluator Python pada lingkungan yang memiliki "
        "koneksi DB dan dependensi (PyMySQL, pandas, NumPy), catat Precision@K, "
        "Recall@K, F1@K, HitRate@K, CatalogCov@K, dan users_evaluated, lalu "
        "lampirkan tanggal eksekusi, hash/versi commit, serta ringkasan "
        "snapshot (jumlah transaksi valid, jumlah user, jumlah produk).",
    )
    add_table(
        doc,
        ["Metadata eksekusi", "Nilai"],
        [
            ["Tanggal/waktu evaluasi", "…"],
            ["Commit / versi aplikasi", "…"],
            ["Jumlah transaksi valid (Selesai+Dibayar)", "…"],
            ["Jumlah pelanggan pada interaksi", "…"],
            ["Jumlah produk pada interaksi", "…"],
            ["Users dievaluasi (holdout)", "…"],
            ["min_co_occurrence", "2"],
            ["Durasi (detik)", "…"],
        ],
        caption="Tabel 9. Metadata reproduksi evaluasi offline",
    )

    add_heading_custom(doc, "4.4 Interpretasi yang Diharapkan setelah Angka Tersedia", 2)
    add_para(
        doc,
        "Karena holdout memakai satu item per pelanggan, Precision@K cenderung "
        "kecil pada katalog besar (paling banyak 1/K bila terjadi hit). "
        "Recall@K sama dengan Hit Rate@K pada skema ini. Catalog Coverage "
        "yang tinggi menunjukkan diversifikasi rekomendasi antar pelanggan, "
        "tetapi bukan bukti akurasi. Klaim keunggulan terhadap baseline lain "
        "hanya sah jika baseline diimplementasikan, dijalankan pada split "
        "yang sama, dan diuji secara statistik dengan metode yang dilampirkan.",
    )

    add_heading_custom(doc, "4.5 Hasil Pengujian Otomatis", 2)
    add_para(
        doc,
        "Arsip proyek memuat 29 metode tes otomatis (13 Python, 16 PHP/Laravel). "
        "Pada lingkungan penyusunan dokumen sebelumnya, sebagian suite Python "
        "(cosine dan parity) dapat dijalankan, sementara evaluator/pipeline "
        "membutuhkan PyMySQL, dan suite PHP membutuhkan binary PHP serta "
        "direktori vendor. Status di bawah harus diperbarui setelah eksekusi "
        "pada lingkungan lengkap.",
    )
    add_table(
        doc,
        ["Suite", "Jumlah tes", "Fokus", "Status target setelah CI lokal"],
        [
            ["Python: test_cosine.py", "5", "Cosine, simetri, threshold, no max-norm", "Wajib lulus"],
            ["Python: test_parity_fixture.py", "2", "Parity fixture kecil vs expected PHP", "Wajib lulus"],
            ["Python: test_evaluator.py", "3", "Leave-last-item, Top-K, anti-leakage", "Wajib lulus"],
            ["Python: test_pipeline.py", "3", "Resolver user, grouping, kolom wajib", "Wajib lulus"],
            ["PHP: CosineSimilarityTest", "6", "Unit cosine PHP", "Wajib lulus"],
            ["PHP: RecommenderServiceTest", "7", "Matrix, CF, fallback, auth admin, dirty flag", "Wajib lulus"],
            ["PHP: ReviewTest", "1", "Ulasan tunggal & anti-duplikasi", "Wajib lulus"],
            ["PHP: ExampleTest (bawaan)", "2", "Scaffold Laravel", "Bukan bukti domain"],
        ],
        caption="Tabel 10. Inventaris pengujian otomatis dan status target",
    )
    add_para(
        doc,
        "Perintah yang disarankan: "
        "python -m unittest discover -s python/tests -v ; "
        "php artisan test --filter=CosineSimilarityTest|RecommenderServiceTest|ReviewTest. "
        "Lampirkan cuplikan keluaran lulus pada lampiran skripsi setelah dijalankan.",
    )

    add_heading_custom(doc, "4.6 Status Pengujian Black-Box dan Integrasi", 2)
    add_para(
        doc,
        "Skenario BB-01 sampai BB-14 telah disusun dari implementasi aktual. "
        "Hasil lulus/gagal hanya boleh diisi setelah aplikasi dijalankan "
        "dengan basis data uji, session, dan (untuk BB-14) Midtrans sandbox. "
        "Inspeksi source code tidak digolongkan sebagai black-box lulus.",
    )
    add_table(
        doc,
        ["Kelompok", "ID", "Status saat ini", "Bukti yang diperlukan"],
        [
            ["Auth & role", "BB-01–BB-02", "Belum dijalankan", "Request/response guest, pelanggan, admin, akun nonaktif"],
            ["Belanja & checkout", "BB-03–BB-04", "Belum dijalankan", "State keranjang, stok, transaksi, rollback gagal"],
            ["Rekomendasi & ulasan", "BB-05–BB-08", "Sebagian dicakup tes kode", "Halaman/API user berriwayat, cold start, stok habis, duplikasi ulasan"],
            ["Admin & upload", "BB-09–BB-12", "Belum dijalankan", "CRUD, validasi file, hash duplikat, progress pipeline, upload_logs"],
            ["Status & pembayaran", "BB-13–BB-14", "Belum dijalankan", "Transisi status; order sandbox; signature; idempotensi"],
        ],
        caption="Tabel 11. Status bukti pengujian black-box dan integrasi",
    )

    add_heading_custom(doc, "4.7 Status UAT, Usability, dan Nonfungsional", 2)
    add_para(
        doc,
        "Arsip proyek menyediakan matriks keterlacakan dan skenario UAT, "
        "tetapi tidak memuat berita acara, record eksekusi, tangkapan layar, "
        "daftar temuan, atau jawaban SUS. Oleh karena itu penerimaan pengguna "
        "dan skor usability belum dapat dinyatakan.",
    )
    add_table(
        doc,
        ["Pengujian", "Instrumen", "Status", "Keputusan saat ini"],
        [
            ["UAT pelanggan", "Skenario P-01–P-18", "Belum dilaksanakan", "Penerimaan pelanggan belum terbukti"],
            ["UAT administrator", "Skenario A-01–A-14", "Belum dilaksanakan", "Kesiapan operasional admin belum terbukti"],
            ["SUS", "10 item; skor 0–100", "Belum ada respons", "Skor usability tidak dapat dihitung"],
            ["Kompatibilitas UI", "Desktop/mobile", "Belum dilaksanakan", "Responsivitas belum diverifikasi"],
            ["Load test Laravel", "p50, p95, throughput, error rate", "Belum dilaksanakan", "Tidak diganti benchmark evaluator Python"],
            ["Keamanan baseline", "Role, session, CSRF, upload, webhook", "Belum dilaksanakan", "Tidak boleh menyatakan aman dari source inspection saja"],
        ],
        caption="Tabel 12. Status UAT, usability, dan nonfungsional",
    )

    add_heading_custom(doc, "4.8 Batas Bukti dan Pengujian Lanjutan", 2)
    add_table(
        doc,
        ["Aspek", "Status", "Tindakan"],
        [
            ["Isi Tabel 8 dari evaluator produksi", "Belum diisi", "Jalankan python/cf/evaluator.py pada snapshot DB penelitian"],
            ["Parity PHP–Python full snapshot", "Belum diuji", "Bandingkan pasangan, skor, dan Top-K pada data sama"],
            ["Black-box BB-01–BB-14", "Belum dijalankan", "Eksekusi HTTP/UI + bukti"],
            ["UAT + SUS", "Belum dilaksanakan", "Sesi penguji + berita acara + respons SUS"],
            ["CTR / conversion / A/B", "Di luar cakupan offline", "Log impresi + eksperimen online bila diperlukan"],
        ],
        caption="Tabel 13. Batas bukti dan pengujian lanjutan",
    )

    add_heading_custom(doc, "4.9 Pembahasan Kesesuaian dengan Implementasi", 2)
    add_para(
        doc,
        "Implementasi produksi memakai IBCF cosine biner dengan ambang "
        "co-occurrence default 2, agregasi prediction_score = sum(sim)/|H_u|, "
        "dan cascade fallback berlabel bukan CF "
        "(Cold Start / Beli Lagi / Produk Tersedia). Serving rekomendasi "
        "pelanggan melalui RecommenderService (PHP) yang membaca "
        "product_similarity; jalur Python dipakai untuk batch/ETL/evaluasi. "
        "Metode pengujian pada Bab III telah diselaraskan dengan "
        "prepare_holdout_splits (leave-last-item) dan metrik evaluate_at_k, "
        "sehingga klaim evaluasi tidak melampaui kemampuan modul di repositori.",
    )
    add_para(
        doc,
        "Perilaku cold-start telah diuji pada tingkat layanan "
        "(RecommenderServiceTest::test_fallback_method_labels_are_not_cf): "
        "label metode fallback tidak mengklaim collaborative filtering. "
        "Transaksi non-valid (bukan Selesai+Dibayar) dikecualikan dari matriks "
        "dan riwayat pembelian untuk rekomendasi, sesuai filter produksi.",
    )

    add_heading_custom(doc, "4.10 Kesimpulan Pengujian", 2)
    add_para(
        doc,
        "1. Teknik rekomendasi yang diuji adalah Item-Based Collaborative "
        "Filtering dengan Cosine Similarity pada matriks interaksi biner; "
        "rating tidak memengaruhi skor IBCF dan hanya relevan pada jalur fallback.",
    )
    add_para(
        doc,
        "2. Evaluasi offline yang selaras repositori menggunakan "
        "leave-last-item holdout berbasis waktu serta metrik Precision@K, "
        "Recall@K, F1@K, Hit Rate@K, dan Catalog Coverage@K melalui "
        "python/cf/evaluator.py (K = 5 dan 10).",
    )
    add_para(
        doc,
        "3. Validitas rumus cosine dan sejumlah perilaku layanan tercakup "
        "suite otomatis (29 metode). Angka akurasi final pada Tabel 8 diisi "
        "setelah evaluator dijalankan pada snapshot Selesai+Dibayar.",
    )
    add_para(
        doc,
        "4. Skenario black-box, UAT, SUS, load test endpoint, dan keamanan "
        "baseline telah dirancang dari implementasi aktual tetapi belum "
        "dilaksanakan pada lingkungan penyusunan dokumen; sistem belum dapat "
        "dinyatakan diterima pengguna atau aman/performan hanya dari inspeksi kode.",
    )
    add_para(
        doc,
        "5. NDCG, baseline Most Popular, dan uji bootstrap tidak dilaporkan "
        "sebagai hasil modul produksi karena tidak diimplementasikan pada "
        "evaluator repositori. Bila diperlukan untuk pembahasan akademik "
        "lanjutan, skrip dan hasilnya harus dilampirkan secara terpisah.",
    )

    add_heading_custom(doc, "REFERENSI", 1)
    refs = [
        "Bangor, A., Kortum, P. T., & Miller, J. T. (2008). An empirical evaluation of the System Usability Scale. International Journal of Human-Computer Interaction, 24(6), 574–594. https://doi.org/10.1080/10447310802205776",
        "Brooke, J. (1996). SUS: A quick and dirty usability scale. In P. W. Jordan, B. Thomas, B. A. Weerdmeester, & I. L. McClelland (Eds.), Usability evaluation in industry (pp. 189–194). Taylor & Francis.",
        "Herlocker, J. L., Konstan, J. A., Terveen, L. G., & Riedl, J. T. (2004). Evaluating collaborative filtering recommender systems. ACM Transactions on Information Systems, 22(1), 5–53. https://doi.org/10.1145/963770.963772",
        "International Organization for Standardization. (2018). ISO 9241-11:2018 Ergonomics of human-system interaction—Part 11: Usability: Definitions and concepts.",
        "International Organization for Standardization. (2021). ISO/IEC/IEEE 29119-3:2021 Software and systems engineering—Software testing—Part 3: Test documentation.",
        "Sarwar, B., Karypis, G., Konstan, J., & Riedl, J. (2001). Item-based collaborative filtering recommendation algorithms. Proceedings of the 10th International World Wide Web Conference, 285–295. https://doi.org/10.1145/371920.372071",
        "Zangerle, E., & Bauer, C. (2022). Evaluating recommender systems: Survey and framework. ACM Computing Surveys, 55(8), Article 170. https://doi.org/10.1145/3556536",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.25)
        p.paragraph_format.first_line_indent = Cm(-1.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        run = p.add_run(ref)
        set_run_font(run, size=11)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(18)
    r = note.add_run(
        "Catatan revisi: Dokumen ini diselaraskan dengan implementasi pada "
        "repositori Toko Sinar Manis (leave-last-item holdout, filter "
        "Selesai+Dibayar, metrik evaluator produksi, inventaris tes otomatis, "
        "dan status jujur untuk black-box/UAT/SUS). Versi lama disimpan sebagai "
        "BAB_III_dan_BAB_IV_Pengujian_Sistem_Rekomendasi_Zaldi_Lengkap_BACKUP.docx."
    )
    set_run_font(r, size=10)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
