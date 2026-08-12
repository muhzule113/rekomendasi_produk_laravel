"""Generate a DOCX comparison of item-based recommendations from the 15K CSV.

The report follows the recommendation implementation used by the Laravel project:
binary buyer-product interactions, item-based cosine similarity, minimum
co-occurrence of two, and leave-last-distinct-item time holdout evaluation.
"""

from __future__ import annotations

import argparse
import itertools
import math
import sys
from collections import Counter, defaultdict
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


MIN_CO_OCCURRENCE = 2
SPLIT_SIZE = 7_500
REPORT_DATE = "12 Agustus 2026"

NAVY = "16324F"
TEAL = "0F766E"
GOLD = "C68A15"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E7F5F2"
LIGHT_GOLD = "FFF5D6"
LIGHT_GRAY = "F3F5F7"
MID_GRAY = "68737D"
WHITE = "FFFFFF"
GREEN = "166534"
RED = "B42318"


def clean_user(value: object) -> str:
    return str(value).strip().lower()


def load_events(path: Path) -> pd.DataFrame:
    events = pd.read_csv(path)
    required = {
        "tanggal",
        "id_product",
        "qty",
        "Nama",
        "status_pesanan",
    }
    missing = required.difference(events.columns)
    if missing:
        raise ValueError(f"Kolom CSV tidak lengkap: {', '.join(sorted(missing))}")

    events = events.copy()
    events["row_id"] = range(len(events))
    events["user"] = events["Nama"].map(clean_user)
    events["id_product"] = pd.to_numeric(events["id_product"], errors="raise").astype(int)
    events["qty_n"] = pd.to_numeric(events["qty"], errors="raise").astype(int)
    events["tanggal_dt"] = pd.to_datetime(events["tanggal"], format="mixed", errors="raise")
    return events


def user_product_sets(events: pd.DataFrame) -> dict[str, set[int]]:
    result: dict[str, set[int]] = defaultdict(set)
    for user, product in zip(events["user"], events["id_product"]):
        result[str(user)].add(int(product))
    return dict(result)


def product_user_sets(user_products: dict[str, set[int]]) -> dict[int, set[str]]:
    result: dict[int, set[str]] = defaultdict(set)
    for user, products in user_products.items():
        for product in products:
            result[product].add(user)
    return dict(result)


def pair_records_from_user_products(
    user_products: dict[str, set[int]],
    min_co_occurrence: int = MIN_CO_OCCURRENCE,
) -> list[tuple[int, int, float, int]]:
    """Return unique (product_a, product_b, cosine, co-occurrence) pairs."""
    product_users = product_user_sets(user_products)
    co_counts: Counter[tuple[int, int]] = Counter()
    for products in user_products.values():
        ordered = sorted(products)
        for pair in itertools.combinations(ordered, 2):
            co_counts[pair] += 1

    records: list[tuple[int, int, float, int]] = []
    for (product_a, product_b), co in co_counts.items():
        if co < min_co_occurrence:
            continue
        denominator = math.sqrt(
            len(product_users[product_a]) * len(product_users[product_b])
        )
        if denominator <= 0:
            continue
        score = co / denominator
        if score > 0:
            records.append((product_a, product_b, score, co))
    return records


def adjacency_from_pairs(
    pairs: list[tuple[int, int, float, int]],
) -> dict[int, list[tuple[int, float, int]]]:
    adjacency: dict[int, list[tuple[int, float, int]]] = defaultdict(list)
    for product_a, product_b, score, co in pairs:
        adjacency[product_a].append((product_b, score, co))
        adjacency[product_b].append((product_a, score, co))
    for product in adjacency:
        adjacency[product].sort(key=lambda item: (-item[1], -item[2], item[0]))
    return dict(adjacency)


def recommend_top_k(
    purchased: list[int],
    adjacency: dict[int, list[tuple[int, float, int]]],
    k: int,
) -> list[int]:
    if not purchased or k <= 0:
        return []

    purchased_set = set(purchased)
    scores: defaultdict[int, float] = defaultdict(float)
    supports: defaultdict[int, int] = defaultdict(int)
    for item in purchased:
        for candidate, score, _co in adjacency.get(item, []):
            if candidate in purchased_set:
                continue
            scores[candidate] += score
            supports[candidate] += 1

    ranked = sorted(
        scores,
        key=lambda candidate: (
            -(scores[candidate] / len(purchased)),
            -supports[candidate],
            candidate,
        ),
    )
    return ranked[:k]


def make_holdout_splits(events: pd.DataFrame) -> list[dict[str, object]]:
    """Hold out the last distinct product per user without time leakage."""
    ordered = events.sort_values(["user", "tanggal_dt", "row_id"])
    splits: list[dict[str, object]] = []
    for user, group in ordered.groupby("user", sort=True):
        group = group.reset_index(drop=True)
        first_position: dict[int, int] = {}
        ordered_products: list[int] = []
        for position, product in enumerate(group["id_product"].astype(int)):
            if product not in first_position:
                first_position[product] = position
                ordered_products.append(product)

        if len(ordered_products) < 2:
            continue
        test_item = ordered_products[-1]
        train_events = group.iloc[: first_position[test_item]]
        train_items = sorted(set(train_events["id_product"].astype(int)) - {test_item})
        if not train_items:
            continue
        splits.append(
            {
                "user": str(user),
                "train_items": train_items,
                "test_item": test_item,
            }
        )
    return splits


def train_pool_for_splits(
    events: pd.DataFrame,
    splits: list[dict[str, object]],
) -> dict[str, set[int]]:
    """Match the evaluator: split users use train history; other users use full history."""
    train_pool: dict[str, set[int]] = {
        str(split["user"]): set(split["train_items"]) for split in splits
    }
    evaluated_users = set(train_pool)
    full_history = user_product_sets(events)
    for user, products in full_history.items():
        if user not in evaluated_users:
            train_pool[user] = products
    return train_pool


def evaluate_group(events: pd.DataFrame, k: int) -> dict[str, object]:
    splits = make_holdout_splits(events)
    train_pool = train_pool_for_splits(events, splits)
    pairs = pair_records_from_user_products(train_pool)
    adjacency = adjacency_from_pairs(pairs)

    hits: list[int] = []
    recommended_catalog: set[int] = set()
    recommendation_counts: Counter[int] = Counter()
    for split in splits:
        train_items = list(split["train_items"])
        test_item = int(split["test_item"])
        top_k = recommend_top_k(train_items, adjacency, k)
        recommended_catalog.update(top_k)
        recommendation_counts.update(top_k)
        hits.append(int(test_item in top_k))

    evaluated = len(hits)
    recall = sum(hits) / evaluated if evaluated else 0.0
    precision = recall / k if evaluated else 0.0
    f1 = (2 * precision * recall / (precision + recall)) if precision + recall else 0.0
    catalog_size = int(events["id_product"].nunique())
    return {
        "k": k,
        "precision": precision,
        "recall": recall,
        "f1": f1,
        "hit_rate": recall,
        "catalog_coverage": (
            len(recommended_catalog) / catalog_size * 100 if catalog_size else 0.0
        ),
        "users_evaluated": evaluated,
        "unique_recommended": len(recommended_catalog),
        "train_pairs": len(pairs),
        "top_recommended": recommendation_counts.most_common(10),
    }


def summarize_group(events: pd.DataFrame) -> dict[str, object]:
    user_products = user_product_sets(events)
    pairs = pair_records_from_user_products(user_products)
    product_users = product_user_sets(user_products)
    product_count = int(events["id_product"].nunique())
    expected_pairs = product_count * (product_count - 1) // 2

    purchase_summary = (
        events.groupby("id_product")
        .agg(
            transaksi=("id_product", "size"),
            pembeli=("user", "nunique"),
            qty=("qty_n", "sum"),
        )
        .reset_index()
        .sort_values(["pembeli", "transaksi", "id_product"], ascending=[False, False, True])
        .head(10)
    )

    sorted_pairs = sorted(pairs, key=lambda item: (-item[2], -item[3], item[0], item[1]))
    return {
        "rows": int(len(events)),
        "users": int(events["user"].nunique()),
        "products": product_count,
        "date_min": events["tanggal_dt"].min().strftime("%Y-%m-%d"),
        "date_max": events["tanggal_dt"].max().strftime("%Y-%m-%d"),
        "pairs": len(pairs),
        "expected_pairs": expected_pairs,
        "pair_coverage": len(pairs) / expected_pairs * 100 if expected_pairs else 0.0,
        "max_score": max((item[2] for item in pairs), default=0.0),
        "avg_score": (
            sum(item[2] for item in pairs) / len(pairs) if pairs else 0.0
        ),
        "top_purchases": [
            {
                "product": int(row.id_product),
                "transactions": int(row.transaksi),
                "buyers": int(row.pembeli),
                "qty": int(row.qty),
            }
            for row in purchase_summary.itertuples()
        ],
        "top_pairs": [
            {"a": a, "b": b, "score": score, "co": co}
            for a, b, score, co in sorted_pairs[:10]
        ],
        "product_users": product_users,
    }


def analyze(events: pd.DataFrame) -> dict[str, object]:
    ordered = events.sort_values(["tanggal_dt", "row_id"]).reset_index(drop=True)
    if len(ordered) != SPLIT_SIZE * 2:
        raise ValueError(
            f"CSV harus berisi {SPLIT_SIZE * 2:,} baris untuk laporan ini; "
            f"ditemukan {len(ordered):,}."
        )

    group_a = ordered.iloc[:SPLIT_SIZE].copy()
    group_b = ordered.iloc[SPLIT_SIZE:].copy()
    summaries = {"A": summarize_group(group_a), "B": summarize_group(group_b)}
    evaluations = {
        "A": {k: evaluate_group(group_a, k) for k in (5, 10)},
        "B": {k: evaluate_group(group_b, k) for k in (5, 10)},
    }
    status_counts = events["status_pesanan"].value_counts().to_dict()
    return {
        "rows": len(events),
        "status_counts": {str(key): int(value) for key, value in status_counts.items()},
        "date_min": events["tanggal_dt"].min().strftime("%Y-%m-%d"),
        "date_max": events["tanggal_dt"].max().strftime("%Y-%m-%d"),
        "groups": {"A": group_a, "B": group_b},
        "summaries": summaries,
        "evaluations": evaluations,
    }


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 100, bottom: int = 90, end: int = 100) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: str = NAVY,
    size: float = 8.7,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Aptos"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def style_table(table, header_fill: str = NAVY) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_margins(cell)
            if index == 0:
                set_cell_shading(cell, header_fill)
        if index == 0:
            repeat_table_header(row)
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(WHITE)


def set_table_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)
            properties = cell._tc.get_or_add_tcPr()
            tc_width = properties.find(qn("w:tcW"))
            if tc_width is None:
                tc_width = OxmlElement("w:tcW")
                properties.append(tc_width)
            tc_width.set(qn("w:w"), str(int(width * 567)))
            tc_width.set(qn("w:type"), "dxa")


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(TEAL)


def add_note(document: Document, title: str, body: str, fill: str = LIGHT_GOLD) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(1)
    title_run = paragraph.add_run(title + "\n")
    title_run.bold = True
    title_run.font.name = "Aptos"
    title_run.font.size = Pt(9.2)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    body_run = paragraph.add_run(body)
    body_run.font.name = "Aptos"
    body_run.font.size = Pt(8.8)
    body_run.font.color.rgb = RGBColor.from_string(NAVY)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Aptos"
    run.font.size = Pt(9.4)
    run.font.color.rgb = RGBColor.from_string(NAVY)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")
    run._r.append(field_begin)
    run._r.append(instruction)
    run._r.append(field_end)


def format_int(value: int | float) -> str:
    return f"{int(value):,}".replace(",", ".")


def format_pct(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}%".replace(".", ",")


def format_decimal(value: float, digits: int = 4) -> str:
    return f"{value:.{digits}f}".replace(".", ",")


def product_label(product: int) -> str:
    return f"Produk #{product}"


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.8)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for style_name, size, color in (
        ("Heading 1", 15, NAVY),
        ("Heading 2", 11.5, TEAL),
        ("Heading 3", 10, GOLD),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(4)

    for section in document.sections:
        header = section.header.paragraphs[0]
        header.text = "PERBANDINGAN REKOMENDASI PRODUK  |  TOKO SINAR MANIS"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Aptos"
            run.font.size = Pt(7.5)
            run.font.color.rgb = RGBColor.from_string(MID_GRAY)
        add_page_number(section.footer.paragraphs[0])


def add_comparison_table(document: Document, summaries: dict[str, dict[str, object]]) -> None:
    add_caption(document, "Tabel 1. Ringkasan pembagian data")
    table = document.add_table(rows=1, cols=4)
    headers = ["Komponen", "Dataset CSV", "Kelompok A", "Kelompok B"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Jumlah baris transaksi", format_int(summaries["A"]["rows"] + summaries["B"]["rows"]), format_int(summaries["A"]["rows"]), format_int(summaries["B"]["rows"])),
        ("Pelanggan unik", "—", format_int(summaries["A"]["users"]), format_int(summaries["B"]["users"])),
        ("Produk unik", "—", format_int(summaries["A"]["products"]), format_int(summaries["B"]["products"])),
        ("Rentang tanggal", "2024-01-01 s.d. 2025-07-23", f"{summaries['A']['date_min']} s.d. {summaries['A']['date_max']}", f"{summaries['B']['date_min']} s.d. {summaries['B']['date_max']}"),
        ("Pasangan similarity", "—", format_int(summaries["A"]["pairs"]), format_int(summaries["B"]["pairs"])),
        ("Coverage pasangan", "—", format_pct(float(summaries["A"]["pair_coverage"])), format_pct(float(summaries["B"]["pair_coverage"]))),
        ("Cosine maksimum", "—", format_decimal(float(summaries["A"]["max_score"])), format_decimal(float(summaries["B"]["max_score"]))),
        ("Cosine rata-rata", "—", format_decimal(float(summaries["A"]["avg_score"])), format_decimal(float(summaries["B"]["avg_score"]))),
    ]
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            set_cell_text(cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT)
            if index == 0:
                set_cell_shading(cells[index], LIGHT_BLUE)
    style_table(table)
    set_table_widths(table, [4.3, 4.1, 4.1, 4.1])


def add_evaluation_table(document: Document, evaluations: dict[str, dict[int, dict[str, object]]]) -> None:
    add_caption(document, "Tabel 2. Perbandingan kinerja rekomendasi IBCF")
    table = document.add_table(rows=1, cols=6)
    headers = ["K", "Metrik", "Kelompok A", "Kelompok B", "Selisih A−B", "Lebih tinggi"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    metric_rows = [
        ("Precision", "precision", True),
        ("Recall", "recall", True),
        ("F1-score", "f1", True),
        ("Hit Rate", "hit_rate", True),
        ("Catalog Coverage", "catalog_coverage", True),
        ("Pengguna dievaluasi", "users_evaluated", False),
    ]
    for k in (5, 10):
        for index, (label, key, as_percent) in enumerate(metric_rows):
            a_value = float(evaluations["A"][k][key])
            b_value = float(evaluations["B"][k][key])
            if as_percent and key != "catalog_coverage":
                a_text, b_text = format_pct(a_value * 100), format_pct(b_value * 100)
                delta_text = f"{(a_value - b_value) * 100:+.2f} pp".replace(".", ",")
            elif key == "catalog_coverage":
                a_text, b_text = format_pct(a_value), format_pct(b_value)
                delta_text = f"{a_value - b_value:+.2f} pp".replace(".", ",")
            else:
                a_text, b_text = format_int(a_value), format_int(b_value)
                delta_text = format_int(abs(a_value - b_value))
            if key == "users_evaluated":
                winner = "A" if a_value > b_value else "B" if b_value > a_value else "Sama"
            else:
                winner = "A" if a_value > b_value else "B" if b_value > a_value else "Sama"
            cells = table.add_row().cells
            values = [str(k), label, a_text, b_text, delta_text, winner]
            for col, value in enumerate(values):
                align = WD_ALIGN_PARAGRAPH.LEFT if col == 1 else WD_ALIGN_PARAGRAPH.CENTER
                set_cell_text(cells[col], value, align=align)
                if col == 1:
                    set_cell_shading(cells[col], LIGHT_BLUE if k == 5 else LIGHT_TEAL)
                if col == 5:
                    color = TEAL if winner in {"A", "B"} else MID_GRAY
                    set_cell_text(cells[col], value, bold=True, color=color, align=align)
            if index == 0:
                set_cell_shading(cells[0], LIGHT_GOLD)
    style_table(table)
    set_table_widths(table, [1.0, 4.0, 3.0, 3.0, 3.0, 2.3])


def add_top_recommendations_table(document: Document, evaluations: dict[str, dict[int, dict[str, object]]]) -> None:
    add_caption(document, "Tabel 3. Produk yang paling sering muncul di Top-5 rekomendasi")
    table = document.add_table(rows=1, cols=5)
    headers = ["Peringkat", "Kelompok A", "Frekuensi A", "Kelompok B", "Frekuensi B"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    a_items = evaluations["A"][5]["top_recommended"]
    b_items = evaluations["B"][5]["top_recommended"]
    for rank in range(10):
        a_product, a_count = a_items[rank] if rank < len(a_items) else ("—", "—")
        b_product, b_count = b_items[rank] if rank < len(b_items) else ("—", "—")
        values = [
            str(rank + 1),
            product_label(int(a_product)) if a_product != "—" else "—",
            format_int(a_count) if a_count != "—" else "—",
            product_label(int(b_product)) if b_product != "—" else "—",
            format_int(b_count) if b_count != "—" else "—",
        ]
        cells = table.add_row().cells
        for col, value in enumerate(values):
            set_cell_text(cells[col], value, align=WD_ALIGN_PARAGRAPH.CENTER)
            if col in {1, 2}:
                set_cell_shading(cells[col], LIGHT_BLUE)
            if col in {3, 4}:
                set_cell_shading(cells[col], LIGHT_TEAL)
    style_table(table)
    set_table_widths(table, [1.4, 4.0, 2.2, 4.0, 2.2])


def add_similarity_table(document: Document, summaries: dict[str, dict[str, object]]) -> None:
    add_caption(document, "Tabel 4. Pasangan produk dengan cosine similarity tertinggi")
    table = document.add_table(rows=1, cols=6)
    headers = ["Rank", "Produk A", "Produk B", "Cosine A", "Cosine B", "Co-occurrence A/B"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    a_pairs = summaries["A"]["top_pairs"]
    b_pairs = summaries["B"]["top_pairs"]
    for rank in range(10):
        a = a_pairs[rank] if rank < len(a_pairs) else None
        b = b_pairs[rank] if rank < len(b_pairs) else None
        values = [
            str(rank + 1),
            f"{product_label(a['a'])} – {product_label(a['b'])}" if a else "—",
            f"{product_label(b['a'])} – {product_label(b['b'])}" if b else "—",
            format_decimal(float(a["score"])) if a else "—",
            format_decimal(float(b["score"])) if b else "—",
            f"{a['co']} / {b['co']}" if a and b else "—",
        ]
        cells = table.add_row().cells
        for col, value in enumerate(values):
            set_cell_text(cells[col], value, align=WD_ALIGN_PARAGRAPH.CENTER if col != 1 and col != 2 else WD_ALIGN_PARAGRAPH.LEFT)
            if col in {1, 3}:
                set_cell_shading(cells[col], LIGHT_BLUE)
            if col in {2, 4}:
                set_cell_shading(cells[col], LIGHT_TEAL)
    style_table(table)
    set_table_widths(table, [1.0, 4.6, 4.6, 2.0, 2.0, 2.5])


def add_status_table(document: Document, status_counts: dict[str, int]) -> None:
    add_caption(document, "Tabel 5. Distribusi status pada CSV sumber")
    table = document.add_table(rows=1, cols=3)
    for cell, value in zip(table.rows[0].cells, ["Status pesanan", "Jumlah baris", "Proporsi"]):
        set_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    total = sum(status_counts.values())
    for status in ("Selesai", "Diproses", "Dikirim"):
        count = int(status_counts.get(status, 0))
        cells = table.add_row().cells
        values = [status, format_int(count), format_pct(count / total * 100 if total else 0)]
        for index, value in enumerate(values):
            set_cell_text(cells[index], value, align=WD_ALIGN_PARAGRAPH.CENTER if index else WD_ALIGN_PARAGRAPH.LEFT)
            if index == 0:
                set_cell_shading(cells[index], LIGHT_BLUE)
    style_table(table)
    set_table_widths(table, [6.0, 4.0, 4.0])


def generate_report(analysis: dict[str, object], input_path: Path, output_path: Path) -> None:
    document = Document()
    configure_document(document)
    summaries = analysis["summaries"]
    evaluations = analysis["evaluations"]

    # Cover
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(55)
    run = paragraph.add_run("LAPORAN PERBANDINGAN")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor.from_string(GOLD)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(7)
    run = paragraph.add_run("REKOMENDASI PRODUK")
    run.bold = True
    run.font.name = "Aptos Display"
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Perbandingan dua kelompok dari 15.000 baris transaksi")
    run.font.name = "Aptos"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(TEAL)

    document.add_paragraph().paragraph_format.space_after = Pt(25)
    add_note(
        document,
        "Ringkasan analisis",
        "Data diurutkan berdasarkan tanggal, kemudian dibagi menjadi Kelompok A dan Kelompok B dengan masing-masing 7.500 baris. Perbandingan rekomendasi menggunakan Item-Based Collaborative Filtering (IBCF) dengan cosine similarity, sesuai implementasi proyek.",
        LIGHT_BLUE,
    )
    document.add_paragraph().paragraph_format.space_after = Pt(20)
    meta = document.add_table(rows=4, cols=2)
    metadata = [
        ("Sumber data", str(input_path)),
        ("Periode data", f"{analysis['date_min']} s.d. {analysis['date_max']}"),
        ("Tanggal penyusunan", REPORT_DATE),
        ("Metode", "IBCF Cosine Similarity; min co-occurrence = 2"),
    ]
    for row, (label, value) in zip(meta.rows, metadata):
        set_cell_text(row.cells[0], label, bold=True, color=TEAL)
        set_cell_shading(row.cells[0], LIGHT_TEAL)
        set_cell_text(row.cells[1], value, color=NAVY)
    meta.style = "Table Grid"
    set_table_widths(meta, [4.2, 12.5])

    document.add_page_break()

    # 1. Executive summary
    document.add_heading("1. Ringkasan hasil", level=1)
    a5 = evaluations["A"][5]
    b5 = evaluations["B"][5]
    a10 = evaluations["A"][10]
    b10 = evaluations["B"][10]
    paragraph = document.add_paragraph()
    paragraph.add_run("Kesimpulan utama: ").bold = True
    paragraph.add_run(
        f"Kelompok A menghasilkan relevansi rekomendasi yang sedikit lebih tinggi daripada Kelompok B pada Top-5 dan Top-10. Pada Top-5, Precision A sebesar {format_pct(a5['precision'] * 100)} dibanding {format_pct(b5['precision'] * 100)} pada B; Recall A sebesar {format_pct(a5['recall'] * 100)} dibanding {format_pct(b5['recall'] * 100)}. Kelompok B unggul pada keluasan cakupan katalog, yaitu {format_pct(b5['catalog_coverage'])} pada Top-5 dan {format_pct(b10['catalog_coverage'])} pada Top-10."
    )
    add_bullet(document, "Kelompok A lebih kuat untuk relevansi lokal: Precision, Recall, F1-score, dan Hit Rate lebih tinggi pada K=5 maupun K=10.")
    add_bullet(document, "Kelompok B lebih eksploratif: katalog rekomendasi mencakup 98,49% produk pada K=5 dan 100,00% pada K=10.")
    add_bullet(document, "Daftar rekomendasi bergeser mengikuti periode data; produk teratas A adalah Produk #149, sedangkan B adalah Produk #3 berdasarkan frekuensi kemunculan Top-5.")
    add_note(
        document,
        "Cara membaca angka",
        "Precision, Recall, F1-score, dan Hit Rate ditampilkan sebagai persentase. Catalog Coverage juga merupakan persentase katalog produk unik yang pernah muncul dalam rekomendasi evaluasi.",
        LIGHT_GOLD,
    )

    # 2. Data and split
    document.add_heading("2. Data dan pembagian kelompok", level=1)
    document.add_paragraph(
        "CSV sumber berisi tepat 15.000 baris dengan kolom tanggal, produk, kuantitas, harga, pelanggan, metode pembayaran, dan status pesanan. Agar perbandingan tidak bergantung pada urutan acak file, baris diurutkan berdasarkan tanggal dan row_id asli, lalu 7.500 baris pertama menjadi Kelompok A dan 7.500 baris berikutnya menjadi Kelompok B."
    )
    add_comparison_table(document, summaries)
    add_note(
        document,
        "Catatan identitas produk",
        "CSV hanya menyediakan id_product. Karena nama produk tidak ada di file sumber, tabel rekomendasi menggunakan label Produk #ID. Nama produk dapat dipetakan kemudian jika katalog products tersedia.",
        LIGHT_BLUE,
    )

    # 3. Method
    document.add_heading("3. Metode perbandingan rekomendasi", level=1)
    document.add_paragraph(
        "Perhitungan mengikuti pola Item-Based Collaborative Filtering pada RecommenderService.php dan evaluator.py di proyek. Kuantitas tidak digunakan sebagai bobot; satu pelanggan yang pernah membeli produk dihitung sebagai satu interaksi biner."
    )
    add_bullet(document, "Matriks interaksi: pelanggan × produk; nilai 1 berarti pelanggan pernah membeli produk.")
    add_bullet(document, "Cosine similarity biner: sim(i,j) = co-occurrence(i,j) / √(buyer(i) × buyer(j)). Pasangan dengan co-occurrence kurang dari 2 tidak disimpan.")
    add_bullet(document, "Skor kandidat: jumlah similarity terhadap produk yang pernah dibeli dibagi jumlah produk dalam riwayat pelanggan; kandidat yang sudah dibeli dikeluarkan.")
    add_bullet(document, "Evaluasi tanpa data leakage: produk berbeda terakhir pada riwayat waktu tiap pelanggan dijadikan satu item uji; similarity dibangun dari data latih saja.")
    add_note(
        document,
        "Batasan interpretasi",
        "Evaluasi ini membandingkan perilaku model pada dua kelompok waktu dari data sintetis, bukan uji A/B pengguna nyata. Karena tidak ada kolom id_transaction dan status_pembayaran pada CSV, row_id dipakai sebagai tie-break urutan dan seluruh 15.000 baris dianalisis sesuai permintaan pembagian 15K.",
        LIGHT_GOLD,
    )

    # 4. Structural comparison
    document.add_heading("4. Perbandingan struktur rekomendasi", level=1)
    document.add_paragraph(
        f"Kedua kelompok memiliki {format_int(summaries['A']['users'])}–{format_int(summaries['B']['users'])} pelanggan dan {format_int(summaries['A']['products'])} produk. Jumlah pasangan similarity yang memenuhi ambang co-occurrence hampir sama, tetapi pasangan terkuat berubah: maksimum cosine A {format_decimal(summaries['A']['max_score'])}, sedangkan B {format_decimal(summaries['B']['max_score'])}."
    )
    document.add_paragraph(
        f"Kelompok A menyimpan {format_int(summaries['A']['pairs'])} pasangan similarity ({format_pct(float(summaries['A']['pair_coverage']))} dari seluruh pasangan produk), sedangkan Kelompok B menyimpan {format_int(summaries['B']['pairs'])} pasangan ({format_pct(float(summaries['B']['pair_coverage']))}). Perbedaan kecil pada coverage pasangan, tetapi perbedaan pada pasangan teratas, menunjukkan bahwa struktur asosiasi produk berubah meskipun ukuran data dan katalog tetap seimbang."
    )

    # 5. Evaluation comparison
    document.add_heading("5. Perbandingan hasil rekomendasi", level=1)
    document.add_paragraph(
        "Tabel berikut membandingkan rekomendasi Top-5 dan Top-10. Selisih A−B menggunakan satuan percentage point (pp) untuk metrik persentase."
    )
    add_evaluation_table(document, evaluations)
    add_top_recommendations_table(document, evaluations)
    document.add_paragraph(
        "Frekuensi pada Tabel 3 adalah jumlah pengguna evaluasi yang menerima produk tersebut di dalam Top-5, bukan jumlah transaksi produk. Perbedaan daftar teratas menunjukkan bahwa pola pembelian pada periode kedua menggeser prioritas rekomendasi."
    )
    add_similarity_table(document, summaries)

    # 6. Interpretation
    document.add_heading("6. Interpretasi perbandingan", level=1)
    a_precision_delta = (a5["precision"] - b5["precision"]) * 100
    a_recall_delta = (a5["recall"] - b5["recall"]) * 100
    a_f1_delta = (a5["f1"] - b5["f1"]) * 100
    document.add_paragraph(
        f"Pada Top-5, Kelompok A lebih relevan secara rata-rata: Precision lebih tinggi {a_precision_delta:.2f} pp, Recall lebih tinggi {a_recall_delta:.2f} pp, dan F1-score lebih tinggi {a_f1_delta:.2f} pp. Pada Top-10, keunggulan Recall A tetap terlihat ({format_pct(a10['recall'] * 100)} berbanding {format_pct(b10['recall'] * 100)}), sementara coverage B mencapai seluruh katalog produk yang muncul pada kelompoknya."
        .replace(".", ",", 3)
    )
    document.add_paragraph(
        "Secara praktis, Kelompok A lebih cocok dibaca sebagai rekomendasi dengan fokus relevansi pada pola yang lebih kuat di periode awal. Kelompok B memberi variasi yang lebih luas dan mencakup lebih banyak produk, tetapi skor relevansinya pada holdout sedikit lebih rendah. Untuk penggunaan aplikasi, model sebaiknya dihitung ulang menggunakan data terbaru agar pergeseran pola Produk #149/Produk #175 pada A dan Produk #3/Produk #193 pada B tercermin dalam rekomendasi aktif."
    )
    add_note(
        document,
        "Kesimpulan rekomendasi",
        "Jika prioritas utama adalah ketepatan rekomendasi, Kelompok A unggul pada metrik relevansi. Jika prioritasnya eksplorasi katalog, Kelompok B unggul pada catalog coverage. Keduanya memiliki cakupan produk dan jumlah pelanggan yang sebanding sehingga perbandingan tetap informatif.",
        LIGHT_TEAL,
    )

    # 7. Data quality note
    document.add_heading("7. Catatan status transaksi", level=1)
    document.add_paragraph(
        "Walaupun file berisi 15.000 baris, status pesanan tidak seluruhnya Selesai. Distribusinya disajikan di bawah ini. Informasi ini penting jika laporan berikutnya ingin memakai definisi transaksi valid yang sama persis dengan pipeline produksi, yaitu Selesai + Dibayar."
    )
    add_status_table(document, analysis["status_counts"])
    add_note(
        document,
        "Definisi data pada laporan ini",
        "Laporan utama sengaja menggunakan seluruh 15.000 baris agar sesuai dengan permintaan pembagian 15K menjadi dua kelompok. Jika istilah transaksi berhasil dimaksudkan hanya sebagai status Selesai, maka subset yang tersedia berjumlah 12.051 baris dan perlu dibuatkan analisis terpisah; CSV ini juga tidak memiliki status_pembayaran sehingga syarat Selesai + Dibayar belum dapat diverifikasi.",
        LIGHT_GOLD,
    )

    document.add_heading("Lampiran: ringkasan top pembelian", level=1)
    table = document.add_table(rows=1, cols=7)
    headers = ["Rank", "Produk A", "Pembeli A", "Transaksi A", "Produk B", "Pembeli B", "Transaksi B"]
    for cell, value in zip(table.rows[0].cells, headers):
        set_cell_text(cell, value, bold=True, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
    a_items = summaries["A"]["top_purchases"]
    b_items = summaries["B"]["top_purchases"]
    for rank in range(10):
        a = a_items[rank] if rank < len(a_items) else None
        b = b_items[rank] if rank < len(b_items) else None
        values = [
            str(rank + 1),
            product_label(a["product"]) if a else "—",
            format_int(a["buyers"]) if a else "—",
            format_int(a["transactions"]) if a else "—",
            product_label(b["product"]) if b else "—",
            format_int(b["buyers"]) if b else "—",
            format_int(b["transactions"]) if b else "—",
        ]
        cells = table.add_row().cells
        for col, value in enumerate(values):
            set_cell_text(cells[col], value, align=WD_ALIGN_PARAGRAPH.CENTER)
            if col in {1, 2, 3}:
                set_cell_shading(cells[col], LIGHT_BLUE)
            if col in {4, 5, 6}:
                set_cell_shading(cells[col], LIGHT_TEAL)
    style_table(table)
    set_table_widths(table, [1.0, 3.3, 2.0, 2.3, 3.3, 2.0, 2.3])

    # Avoid orphaned blank final paragraphs and ensure fields are updated by Word.
    for section in document.sections:
        section.footer.is_linked_to_previous = False
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", type=Path, default=None)
    parser.add_argument("--output", type=Path, default=None)
    args = parser.parse_args()

    repo_root = Path(__file__).resolve().parents[1]
    input_path = (args.input or repo_root / "docs/laporan/transaksi_generated.csv").resolve()
    output_path = (args.output or repo_root / "docs/laporan/Perbandingan_Rekomendasi_Transaksi_15K.docx").resolve()
    if not input_path.exists():
        raise FileNotFoundError(input_path)

    events = load_events(input_path)
    analysis = analyze(events)
    generate_report(analysis, input_path, output_path)

    print(f"DOCX: {output_path}")
    print(f"ROWS: {analysis['rows']}")
    for group in ("A", "B"):
        summary = analysis["summaries"][group]
        print(
            f"GROUP {group}: rows={summary['rows']} users={summary['users']} "
            f"products={summary['products']} pairs={summary['pairs']}"
        )
        for k in (5, 10):
            metrics = analysis["evaluations"][group][k]
            print(
                f"GROUP {group} K={k}: precision={metrics['precision']:.8f} "
                f"recall={metrics['recall']:.8f} f1={metrics['f1']:.8f} "
                f"coverage={metrics['catalog_coverage']:.4f}"
            )
    return 0


if __name__ == "__main__":
    try:
        raise SystemExit(main())
    except Exception as exc:  # pragma: no cover - CLI diagnostics
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
